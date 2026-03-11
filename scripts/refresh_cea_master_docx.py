#!/usr/bin/env python
"""Refresh the CEA Word manuscript using a user-edited DOCX as text source."""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


IMAGE_CAPTION_STYLE = "Image Caption"


@dataclass
class ImageInfo:
    caption: str
    rid: str
    blob: bytes
    wp_extent: tuple[str, str] | None
    a_ext: tuple[str, str] | None


def normalize_text(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def para_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath('.//*[local-name()="drawing"]'))


def first_embed_rid(paragraph: Paragraph) -> str | None:
    blips = paragraph._p.xpath('.//*[local-name()="blip"]')
    for blip in blips:
        rid = blip.get(qn("r:embed"))
        if rid:
            return rid
    return None


def get_extent(paragraph: Paragraph, local_name: str) -> tuple[str, str] | None:
    nodes = paragraph._p.xpath(f'.//*[local-name()="{local_name}"]')
    for node in nodes:
        cx = node.get("cx")
        cy = node.get("cy")
        if cx is not None and cy is not None:
            return cx, cy
    return None


def set_extent(paragraph: Paragraph, local_name: str, extent: tuple[str, str] | None) -> None:
    if extent is None:
        return
    cx, cy = extent
    nodes = paragraph._p.xpath(f'.//*[local-name()="{local_name}"]')
    for node in nodes:
        if node.get("cx") is not None:
            node.set("cx", cx)
        if node.get("cy") is not None:
            node.set("cy", cy)


def find_previous_image_paragraph(paragraphs: list[Paragraph], start_idx: int) -> Paragraph | None:
    for idx in range(start_idx - 1, -1, -1):
        if para_has_drawing(paragraphs[idx]):
            return paragraphs[idx]
        text = paragraphs[idx].text.strip()
        if text:
            break
    return None


def collect_images(doc: DocumentType) -> dict[str, ImageInfo]:
    images: dict[str, ImageInfo] = {}
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        if not paragraph.style or paragraph.style.name != IMAGE_CAPTION_STYLE:
            continue
        caption = normalize_text(paragraph.text)
        if not caption:
            continue
        image_paragraph = find_previous_image_paragraph(paragraphs, idx)
        if image_paragraph is None:
            continue
        rid = first_embed_rid(image_paragraph)
        if rid is None:
            continue
        part = doc.part.related_parts[rid]
        images[caption] = ImageInfo(
            caption=caption,
            rid=rid,
            blob=part.blob,
            wp_extent=get_extent(image_paragraph, "extent"),
            a_ext=get_extent(image_paragraph, "ext"),
        )
    return images


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    element = paragraph._element
    for child in list(element):
        element.remove(child)
    paragraph.add_run(text)


def replace_substrings(paragraph: Paragraph, replacements: Iterable[tuple[str, str]]) -> bool:
    text = paragraph.text
    new_text = text
    for old, new in replacements:
        new_text = new_text.replace(old, new)
    if new_text != text:
        set_paragraph_text(paragraph, new_text)
        return True
    return False


def apply_global_text_fixes(doc: DocumentType) -> list[str]:
    replacements = [
        ("Byte Track", "ByteTrack"),
        ("byteTrack", "ByteTrack"),
        ("Byte-track", "ByteTrack"),
    ]
    change_log: list[str] = []
    for idx, paragraph in enumerate(doc.paragraphs, 1):
        if replace_substrings(paragraph, replacements):
            change_log.append(f"Normalized terminology at paragraph {idx}.")

        text = paragraph.text
        normalized = normalize_text(text)
        if "C3: Deployment decision protocol" in normalized and "Section 6" in normalized:
            set_paragraph_text(paragraph, normalized.replace("Section 6", "Section 5"))
            change_log.append(f"Fixed C3 section reference at paragraph {idx}: Section 6 -> Section 5.")
            continue

        if "activation analysis in Section 5.5" in normalized:
            set_paragraph_text(paragraph, normalized.replace("Section 5.5", "Section 4.5"))
            change_log.append(f"Fixed gating activation reference at paragraph {idx}: Section 5.5 -> Section 4.5.")
            continue

        if "Section 5.5 specifies" in normalized:
            set_paragraph_text(paragraph, normalized.replace("Section 5.5", "Section 4.5"))
            change_log.append(f"Fixed discussion back-reference at paragraph {idx}: Section 5.5 -> Section 4.5.")
            continue

        if "report all fields in Table 3" in normalized:
            set_paragraph_text(paragraph, normalized.replace("Table 3", "Table 7"))
            change_log.append(f"Fixed runtime schema table reference at paragraph {idx}: Table 3 -> Table 7.")

    return change_log


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_duplicate_lines(doc: DocumentType) -> list[str]:
    change_log: list[str] = []
    previous_text = None
    for paragraph in list(doc.paragraphs):
        text = normalize_text(paragraph.text)
        if not text:
            previous_text = None
            continue
        if text == previous_text:
            remove_paragraph(paragraph)
            change_log.append(f"Removed duplicate paragraph: {text[:80]}")
            continue
        previous_text = text
    return change_log


def build_submission_copy(in_path: Path, out_path: Path) -> list[str]:
    doc = Document(in_path)
    change_log: list[str] = []
    author_contrib_replacement = (
        "Conceptualization, Methodology, Software, Validation, Formal analysis, "
        "Investigation, Data curation, Writing – original draft, Writing – review & editing, "
        "Visualization: author team. Supervision, Project administration, Resources: author team. "
        "All authors read and approved the final manuscript."
    )

    for paragraph in list(doc.paragraphs):
        text = normalize_text(paragraph.text)
        style = paragraph.style.name if paragraph.style else ""
        if style == "Author":
            remove_paragraph(paragraph)
            change_log.append("Removed blinded author line from submission copy.")
            continue
        if text.startswith("Title page information (to be completed in the final submission metadata):"):
            remove_paragraph(paragraph)
            change_log.append("Removed title-page placeholder block from submission copy.")
            continue
        if text.startswith("Desk-reject prevention (CEA):"):
            remove_paragraph(paragraph)
            change_log.append("Removed author-only desk-reject checklist from submission copy.")
            continue
        if text == "Appendix B. Pre-submission checklist (author-only; remove before final submission)":
            remove_paragraph(paragraph)
            change_log.append("Removed author-only pre-submission checklist heading.")
            continue
        if style == "First Paragraph" and text.startswith("Conceptualization, Methodology, Software"):
            set_paragraph_text(paragraph, author_contrib_replacement)
            change_log.append("Replaced author-contribution placeholders with blinded wording.")

    change_log.extend(remove_duplicate_lines(doc))
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return change_log


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Refresh the CEA Word manuscript from a master DOCX.")
    parser.add_argument("--master-docx", required=True, type=Path, help="User-edited DOCX used as text source.")
    parser.add_argument("--donor-docx", required=True, type=Path, help="DOCX containing the figure assets to reuse.")
    parser.add_argument("--master-out", required=True, type=Path, help="Output DOCX path for the refreshed master file.")
    parser.add_argument(
        "--submission-out",
        default=None,
        type=Path,
        help="Optional output DOCX path for a cleaned submission copy.",
    )
    parser.add_argument("--report", default=None, type=Path, help="Optional markdown/text report path.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    master_doc = Document(args.master_docx)
    donor_doc = Document(args.donor_docx)

    donor_images = collect_images(donor_doc)
    master_images = collect_images(master_doc)

    change_log: list[str] = []
    for caption, master_info in master_images.items():
        donor_info = donor_images.get(caption)
        if donor_info is None:
            change_log.append(f"Skipped figure replacement for caption without donor match: {caption}")
            continue
        master_part = master_doc.part.related_parts[master_info.rid]
        master_part._blob = donor_info.blob
        paragraphs = list(master_doc.paragraphs)
        for idx, paragraph in enumerate(paragraphs):
            if paragraph.style and paragraph.style.name == IMAGE_CAPTION_STYLE and normalize_text(paragraph.text) == caption:
                image_paragraph = find_previous_image_paragraph(paragraphs, idx)
                if image_paragraph is not None:
                    set_extent(image_paragraph, "extent", donor_info.wp_extent)
                    set_extent(image_paragraph, "ext", donor_info.a_ext)
                break
        change_log.append(f"Replaced figure asset for caption: {caption}")

    change_log.extend(apply_global_text_fixes(master_doc))

    args.master_out.parent.mkdir(parents=True, exist_ok=True)
    master_doc.save(args.master_out)

    if args.submission_out is not None:
        change_log.extend(build_submission_copy(args.master_out, args.submission_out))

    if args.report is not None:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text("\n".join(f"- {line}" for line in change_log) + "\n", encoding="utf-8")

    print(f"Saved refreshed master DOCX: {args.master_out}")
    if args.submission_out is not None:
        print(f"Saved submission DOCX: {args.submission_out}")
    if args.report is not None:
        print(f"Saved report: {args.report}")
    print(f"figure_matches={len(master_images)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
