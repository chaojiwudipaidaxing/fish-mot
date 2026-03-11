from __future__ import annotations

import argparse
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}


def doc_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def check(condition: bool, name: str, detail: str, lines: list[str]) -> int:
    lines.append(f"[{'PASS' if condition else 'FAIL'}] {name}: {detail}")
    return 0 if condition else 1


def main() -> int:
    ap = argparse.ArgumentParser(description="Validate strict COMPAG outputs.")
    ap.add_argument("--outdir", default="out_compag_strict")
    args = ap.parse_args()

    outdir = Path(args.outdir)
    en_docx = outdir / "COMPAG_CEA_manuscript_EN.docx"
    en_pdf = outdir / "COMPAG_CEA_manuscript_EN.pdf"
    zh_docx = outdir / "COMPAG_CEA_manuscript_ZH_working.docx"
    zh_pdf = outdir / "COMPAG_CEA_manuscript_ZH_working.pdf"
    highlights = outdir / "highlights.docx"
    ga = outdir / "graphical_abstract_placeholder.docx"
    changes = outdir / "changes.md"
    todo = outdir / "todo.md"

    lines: list[str] = []
    fail = 0
    for p in [en_docx, en_pdf, zh_docx, zh_pdf, highlights, ga, changes, todo]:
        fail += check(p.exists(), "file exists", str(p), lines)

    if not en_docx.exists():
        report = outdir / "validation_report.txt"
        report.write_text("\n".join(lines) + "\n", encoding="utf-8")
        return 1

    d = Document(en_docx)
    txt = doc_text(d)
    # Basic compliance
    abs_words = 0
    key_count = 0
    paras = [p.text.strip() for p in d.paragraphs]
    for i, t in enumerate(paras):
        if t.lower() == "abstract" and i + 1 < len(paras):
            abs_words = len(re.findall(r"[A-Za-z0-9\-]+", paras[i + 1]))
            if i + 2 < len(paras):
                key_count = len([x.strip() for x in paras[i + 2].split(",") if x.strip()])
            break
    fail += check(abs_words <= 250, "abstract <=250", f"words={abs_words}", lines)
    fail += check(1 <= key_count <= 7, "keywords 1..7", f"count={key_count}", lines)

    # token cleanup
    for token in ["Table Table", "[tab:", "[alg:", "meanstd", "drift-proof", "zenodo.XXXX", "anonymous.4open", "Let  be"]:
        fail += check(token not in txt, f"token removed `{token}`", f"count={txt.count(token)}", lines)

    # required phrases
    required_phrases = [
        "Highlights: submitted as a separate file.",
        "Graphical abstract: submitted as a separate file",
        "Graphical Abstract - Figure GA (submitted separately).",
        "Data availability.",
        "Terminology and scope",
        "Scope A",
        "Scope B",
        "Gating activation analysis",
    ]
    for phrase in required_phrases:
        fail += check(phrase in txt, f"contains `{phrase}`", "", lines)

    # scope B fields
    scope_fields = {
        "fps_e2e",
        "mem_peak_mb_e2e",
        "cpu_norm_e2e",
        "decode_time",
        "detector_time",
        "tracking_time",
        "write_time",
        "detector_name",
        "input_resolution",
    }
    found_fields = set()
    for t in d.tables:
        for r in t.rows:
            if r.cells:
                found_fields.add(r.cells[0].text.strip())
    fail += check(scope_fields.issubset(found_fields), "Scope B field coverage", f"missing={sorted(scope_fields - found_fields)}", lines)

    # equations and OMML
    with zipfile.ZipFile(en_docx) as z:
        root = ET.fromstring(z.read("word/document.xml"))
    omml_count = len(root.findall(".//m:oMath", NS)) + len(root.findall(".//m:oMathPara", NS))
    fail += check(omml_count > 0, "OMML equation presence", f"count={omml_count}", lines)

    # one-column check
    multi = 0
    for cols in root.findall(".//w:sectPr/w:cols", NS):
        num = cols.attrib.get(f"{{{NS['w']}}}num")
        if num and num != "1":
            multi += 1
    fail += check(multi == 0, "single-column layout", f"multi_sections={multi}", lines)

    # vertical border check: should avoid vertical rules
    bad_vertical = 0
    for edge in root.findall(".//w:tcBorders/w:left", NS) + root.findall(".//w:tcBorders/w:right", NS):
        val = edge.attrib.get(f"{{{NS['w']}}}val", "")
        if val and val not in {"nil", "none"}:
            bad_vertical += 1
    fail += check(bad_vertical == 0, "no vertical table rules", f"bad_vertical_edges={bad_vertical}", lines)

    # Highlights constraint
    if highlights.exists():
        hd = Document(highlights)
        items = [p.text.strip().lstrip("•").strip() for p in hd.paragraphs if p.text.strip()]
        fail += check(3 <= len(items) <= 5, "highlights count 3..5", f"count={len(items)}", lines)
        fail += check(all(len(x) <= 85 for x in items), "highlights <=85 chars", f"lengths={[len(x) for x in items]}", lines)

    report = outdir / "validation_report.txt"
    lines.append(f"\nsummary: total={len(lines)} fail={fail}")
    report.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print("\n".join(lines))
    return 1 if fail else 0


if __name__ == "__main__":
    raise SystemExit(main())
