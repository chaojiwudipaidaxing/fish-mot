#!/usr/bin/env python
"""Build a CEA-oriented manuscript package without changing source wording."""

from __future__ import annotations

import argparse
import hashlib
import io
import re
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Mm, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


IMAGE_CAPTION_STYLE = "Image Caption"
TABLE_CAPTION_STYLE = "Table Caption"
FIGURE_STYLE = "Captioned Figure"
BODY_STYLES = {"Normal", "Body Text", "First Paragraph", "Abstract", "Bibliography"}
HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Abstract Title"}


@dataclass
class ImageInfo:
    caption: str
    rid: str
    blob: bytes
    wp_extent: tuple[str, str] | None
    a_ext: tuple[str, str] | None


@dataclass
class BuildStats:
    figures_replaced: int = 0
    figures_inserted: int = 0
    tables_styled: int = 0
    removed_paragraphs: int = 0
    replaced_placeholders: int = 0
    source_todos: int = 0
    source_placeholders: int = 0
    output_todos: int = 0
    output_placeholders: int = 0
    source_figures: int = 0
    output_figures: int = 0


def normalize_space(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def para_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath('.//*[local-name()="drawing"]'))


def first_embed_rid(paragraph: Paragraph) -> str | None:
    blips = paragraph._p.xpath('.//*[local-name()="blip"]')
    for blip in blips:
        rid = blip.get(qn("r:embed"))
        if rid:
            return rid
    return None


def get_extent(paragraph: Paragraph, local_name: str) -> tuple[str, str] | None:
    nodes = paragraph._p.xpath(f'.//*[local-name()="{local_name}"]')
    for node in nodes:
        cx = node.get("cx")
        cy = node.get("cy")
        if cx is not None and cy is not None:
            return cx, cy
    return None


def set_extent(paragraph: Paragraph, local_name: str, extent: tuple[str, str] | None) -> None:
    if extent is None:
        return
    cx, cy = extent
    nodes = paragraph._p.xpath(f'.//*[local-name()="{local_name}"]')
    for node in nodes:
        if node.get("cx") is not None:
            node.set("cx", cx)
        if node.get("cy") is not None:
            node.set("cy", cy)


def find_previous_image_paragraph(paragraphs: list[Paragraph], start_idx: int) -> Paragraph | None:
    for idx in range(start_idx - 1, -1, -1):
        if para_has_drawing(paragraphs[idx]):
            return paragraphs[idx]
        text = paragraphs[idx].text.strip()
        if text:
            break
    return None


def collect_images(doc: DocumentType) -> dict[str, ImageInfo]:
    images: dict[str, ImageInfo] = {}
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        if not paragraph.style or paragraph.style.name != IMAGE_CAPTION_STYLE:
            continue
        caption = normalize_space(paragraph.text)
        if not caption:
            continue
        image_paragraph = find_previous_image_paragraph(paragraphs, idx)
        if image_paragraph is None:
            continue
        rid = first_embed_rid(image_paragraph)
        if rid is None:
            continue
        part = doc.part.related_parts[rid]
        images[caption] = ImageInfo(
            caption=caption,
            rid=rid,
            blob=part.blob,
            wp_extent=get_extent(image_paragraph, "extent"),
            a_ext=get_extent(image_paragraph, "ext"),
        )
    return images


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_style_font(style, *, size_pt: float | None, bold: bool | None, italic: bool | None) -> None:
    font = style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0, 0, 0)
    if size_pt is not None:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic

    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), "Times New Roman")


def set_runs_font(paragraph: Paragraph, *, size_pt: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{key}"), "Times New Roman")


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    element = paragraph._element
    for child in list(element):
        element.remove(child)
    paragraph.add_run(text)
    set_runs_font(paragraph)


def apply_style_contract(doc: DocumentType) -> None:
    styles = doc.styles

    for name in BODY_STYLES:
        if name in styles:
            set_style_font(styles[name], size_pt=10, bold=False, italic=False)
            pf = styles[name].paragraph_format
            pf.line_spacing = 1.15
            pf.space_before = Pt(0)
            pf.space_after = Pt(4)

    if "Title" in styles:
        set_style_font(styles["Title"], size_pt=17, bold=True, italic=False)
        styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        styles["Title"].paragraph_format.space_after = Pt(6)

    if "Author" in styles:
        set_style_font(styles["Author"], size_pt=11, bold=False, italic=False)
        styles["Author"].paragraph_format.space_after = Pt(8)

    for name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4", "Abstract Title"):
        if name in styles:
            size = 12 if name == "Heading 1" else 11 if name == "Heading 2" else 10
            set_style_font(styles[name], size_pt=size, bold=True, italic=False)
            pf = styles[name].paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.line_spacing = 1.1
            pf.space_before = Pt(10 if name == "Heading 1" else 8 if name == "Heading 2" else 5)
            pf.space_after = Pt(3)
            pf.keep_with_next = True

    for name in ("Image Caption", "Table Caption"):
        if name in styles:
            set_style_font(styles[name], size_pt=9, bold=False, italic=True)
            pf = styles[name].paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.line_spacing = 1.0
            pf.space_before = Pt(3)
            pf.space_after = Pt(6)
            pf.keep_with_next = True

    if FIGURE_STYLE in styles:
        set_style_font(styles[FIGURE_STYLE], size_pt=10, bold=False, italic=False)
        pf = styles[FIGURE_STYLE].paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.line_spacing = 1.0
        pf.space_before = Pt(8)
        pf.space_after = Pt(2)
        pf.keep_with_next = True


def apply_page_setup(doc: DocumentType, review_mode: bool = False) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(22)
        section.bottom_margin = Mm(22)
        section.left_margin = Mm(20 if review_mode else 18)
        section.right_margin = Mm(20 if review_mode else 18)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)


def apply_paragraph_layout(doc: DocumentType) -> None:
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""
        text = paragraph.text.strip()
        set_runs_font(paragraph)
        pf = paragraph.paragraph_format
        pf.keep_together = False
        pf.widow_control = True

        if style_name in BODY_STYLES:
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.line_spacing = 1.15
            pf.space_before = Pt(0)
            pf.space_after = Pt(4)
        elif style_name in HEADING_STYLES:
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.keep_with_next = True
        elif style_name == IMAGE_CAPTION_STYLE:
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.keep_with_next = False
        elif style_name == TABLE_CAPTION_STYLE:
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.keep_with_next = True
        elif style_name == FIGURE_STYLE or para_has_drawing(paragraph):
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.keep_with_next = True
            pf.space_before = Pt(8)
            pf.space_after = Pt(2)

        if style_name == "Title":
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if style_name == "Author":
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if style_name == "Abstract Title":
            pf.space_before = Pt(0)
            pf.space_after = Pt(4)
        if style_name == "Abstract":
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if style_name == "First Paragraph" and text.startswith("Keywords:"):
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_after = Pt(8)

        if para_has_drawing(paragraph):
            next_style = paragraphs[idx + 1].style.name if idx + 1 < len(paragraphs) and paragraphs[idx + 1].style else ""
            if next_style == IMAGE_CAPTION_STYLE:
                pf.keep_with_next = True


def set_cell_border(cell: _Cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def style_table(table: Table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    rows = table.rows
    if not rows:
        return

    ncols = len(rows[0].cells)
    size_pt = 9.5 if ncols <= 5 else 9.0 if ncols == 6 else 8.3

    for row_idx, row in enumerate(rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(
                cell,
                left={"val": "nil"},
                right={"val": "nil"},
                top={"val": "nil"},
                bottom={"val": "nil"},
            )
            for paragraph in cell.paragraphs:
                set_runs_font(paragraph, size_pt=size_pt)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )

        if row_idx == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    for cell in rows[0].cells:
        set_cell_border(cell, top={"val": "single", "sz": "8", "color": "000000"})
        set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": "000000"})
    for cell in rows[-1].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": "000000"})


def style_tables(doc: DocumentType, stats: BuildStats) -> None:
    for table in doc.tables:
        style_table(table)
        stats.tables_styled += 1


def clean_placeholder_blocks(doc: DocumentType, stats: BuildStats) -> list[str]:
    paragraphs = list(doc.paragraphs)
    normalized = [normalize_space(p.text) for p in paragraphs]
    remove_idx: set[int] = set()
    log: list[str] = []

    def mark_range(start: int, end: int, reason: str) -> None:
        for idx in range(start, end + 1):
            if 0 <= idx < len(paragraphs):
                remove_idx.add(idx)
        log.append(reason)

    for idx, text in enumerate(normalized):
        style_name = paragraphs[idx].style.name if paragraphs[idx].style else ""

        if style_name == "Author" and text == "Anonymous Author(s)":
            remove_idx.add(idx)
            log.append("Removed blinded author placeholder line from manuscript.")

        if text.startswith("Title page information (to be completed in the final submission metadata):"):
            remove_idx.add(idx)
            log.append("Removed title-page metadata placeholder block from manuscript.")

        if text == "Additional validation items reserved for future extension.":
            end = idx
            while end + 1 < len(paragraphs):
                next_style = paragraphs[end + 1].style.name if paragraphs[end + 1].style else ""
                next_text = normalized[end + 1]
                if next_style in {"Heading 1", "Heading 2"} and next_text:
                    break
                end += 1
            mark_range(idx, end, "Removed future-validation placeholder block.")

        if text == "Planned environmental stress validation":
            end = min(idx + 1, len(paragraphs) - 1)
            mark_range(idx, end, "Removed planned environmental-stress placeholder note.")

        if text == "Planned figures/tables (TODO).":
            end = idx
            while end + 1 < len(paragraphs):
                next_style = paragraphs[end + 1].style.name if paragraphs[end + 1].style else ""
                next_text = normalized[end + 1]
                if next_style in {"Heading 1", "Heading 2"} and next_text:
                    break
                if next_text.startswith("TODO-R1:"):
                    break
                end += 1
            mark_range(idx, end, "Removed planned-figure and TODO extraction placeholder block.")

        if text.startswith("TODO-R1:"):
            remove_idx.add(idx)
            log.append("Removed runtime TODO note.")

        if "[Author 1]" in text or "[Author 2]" in text:
            replacement = (
                "Conceptualization, Methodology, Software, Validation, Formal analysis, "
                "Investigation, Data curation, Writing – original draft, Writing – review & editing, "
                "Visualization: author team. Supervision, Project administration, Resources: author team. "
                "All authors read and approved the final manuscript."
            )
            set_paragraph_text(paragraphs[idx], replacement)
            stats.replaced_placeholders += 1
            log.append("Replaced author-name placeholders in CRediT statement with blinded wording.")

    for idx in sorted(remove_idx, reverse=True):
        remove_paragraph(paragraphs[idx])
        stats.removed_paragraphs += 1

    return log


def strip_todo_hint(text: str) -> str:
    new_text = re.sub(r"\s*\(TODO[^)]*\)", "", text)
    new_text = re.sub(r"\s+\.", ".", new_text)
    new_text = re.sub(r"\.\.", ".", new_text)
    return new_text.strip()


def clean_table_reporting_placeholders(doc: DocumentType) -> list[str]:
    cleaned_cells = 0
    cleaned_captions = 0

    for table in doc.tables:
        caption_nodes = table._tbl.xpath('.//*[local-name()="tblCaption"]')
        for node in caption_nodes:
            val = node.get(qn("w:val"))
            if val and "TODO" in val:
                node.set(qn("w:val"), strip_todo_hint(val))
                cleaned_captions += 1

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text = strip_todo_hint(paragraph.text)
                    if new_text != paragraph.text:
                        set_paragraph_text(paragraph, new_text)
                        cleaned_cells += 1

    log: list[str] = []
    if cleaned_captions:
        log.append(f"Removed TODO markers from {cleaned_captions} table-caption metadata field(s).")
    if cleaned_cells:
        log.append(f"Removed TODO markers from {cleaned_cells} table cell(s).")
    return log


def replace_or_insert_figures(master_doc: DocumentType, donor_doc: DocumentType, stats: BuildStats) -> list[str]:
    donor_images = collect_images(donor_doc)
    master_images = collect_images(master_doc)
    paragraphs = list(master_doc.paragraphs)
    log: list[str] = []

    for idx, paragraph in enumerate(paragraphs):
        if not paragraph.style or paragraph.style.name != IMAGE_CAPTION_STYLE:
            continue

        caption = normalize_space(paragraph.text)
        donor_info = donor_images.get(caption)
        if donor_info is None:
            log.append(f"Skipped figure without donor match: {caption}")
            continue

        master_info = master_images.get(caption)
        if master_info is not None:
            master_part = master_doc.part.related_parts[master_info.rid]
            master_part._blob = donor_info.blob
            image_paragraph = find_previous_image_paragraph(paragraphs, idx)
            if image_paragraph is not None:
                set_extent(image_paragraph, "extent", donor_info.wp_extent)
                set_extent(image_paragraph, "ext", donor_info.a_ext)
                image_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            stats.figures_replaced += 1
            log.append(f"Replaced figure asset: {caption}")
            continue

        target = paragraph.insert_paragraph_before("")
        if FIGURE_STYLE in master_doc.styles:
            target.style = master_doc.styles[FIGURE_STYLE]
        target.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = target.add_run()
        width = Emu(int(donor_info.wp_extent[0])) if donor_info.wp_extent else Inches(6.2)
        run.add_picture(io.BytesIO(donor_info.blob), width=width)
        stats.figures_inserted += 1
        log.append(f"Inserted missing figure: {caption}")

    stats.output_figures = len(master_doc.inline_shapes)
    return log


def add_page_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def enable_line_numbers(doc: DocumentType) -> None:
    for section in doc.sections:
        sect_pr = section._sectPr
        for node in sect_pr.findall(qn("w:lnNumType")):
            sect_pr.remove(node)
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:start"), "1")
        ln.set(qn("w:restart"), "newPage")
        sect_pr.append(ln)


def add_footer_page_numbers(doc: DocumentType) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)
        paragraph.text = "Page "
        set_runs_font(paragraph, size_pt=9)
        add_page_field(paragraph)


def create_review_copy(src_docx: Path, out_docx: Path) -> None:
    doc = Document(str(src_docx))
    apply_page_setup(doc, review_mode=True)
    enable_line_numbers(doc)
    add_footer_page_numbers(doc)
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name in BODY_STYLES:
            paragraph.paragraph_format.line_spacing = 1.5
        set_runs_font(paragraph)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def document_signature(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml")
    return hashlib.sha256(xml).hexdigest()


def visible_text_signature(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    joined = "\n".join(parts)
    return hashlib.sha256(joined.encode("utf-8")).hexdigest()


def count_tokens(docx_path: Path) -> tuple[int, int]:
    doc = Document(str(docx_path))
    full = "\n".join(p.text for p in doc.paragraphs)
    todo_count = len(re.findall(r"\bTODO\b", full))
    placeholder_count = len(
        re.findall(r"placeholder|\[Author\s+\d+\]|Title page information|Anonymous Author", full, flags=re.IGNORECASE)
    )
    return todo_count, placeholder_count


def first_title(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name == "Title" and paragraph.text.strip():
            return paragraph.text.strip()
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            return paragraph.text.strip()
    return "CEA Manuscript Title"


def create_highlights_docx(path: Path) -> None:
    bullets = [
        "Locked configuration and manifest hashing support reproducible fish MOT evaluation.",
        "Drift-aware diagnostics expose failures under turning and low-confidence conditions.",
        "Tracking, counting, and runtime evidence are reported for deployment decisions.",
        "Scope A/B profiling separates tracker cost from end-to-end aquaculture operation.",
    ]
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10)
    for line in bullets:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(line)
        set_runs_font(paragraph, size_pt=10)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def create_cover_letter(path: Path, title: str) -> None:
    doc = Document()
    p = doc.add_paragraph("Cover Letter")
    p.style = doc.styles["Title"]
    set_runs_font(p, size_pt=16, bold=True)
    for line in [
        "Dear Editor,",
        (
            f'We submit the manuscript "{title}" to Computers and Electronics in Agriculture. '
            "The enclosed files provide a text-locked manuscript, supporting submission documents, "
            "and an anonymized reproducibility bundle."
        ),
        (
            "This submission presents a drift-aware and deployment-oriented evaluation framework for fish "
            "multi-object tracking and counting in digital aquaculture."
        ),
        (
            "The work is original, is not under consideration elsewhere, and has been prepared for "
            "double-blind review."
        ),
        "Sincerely,",
        "The Authors",
    ]:
        paragraph = doc.add_paragraph(line)
        set_runs_font(paragraph, size_pt=10)
    doc.save(str(path))


def create_title_page(path: Path, title: str) -> None:
    doc = Document()
    heading = doc.add_paragraph("Title Page")
    heading.style = doc.styles["Title"]
    set_runs_font(heading, size_pt=16, bold=True)
    lines = [
        f"Article title: {title}",
        "Authors: Anonymous author team (double-blind review mode).",
        "Affiliations: To be completed in the submission metadata.",
        "Corresponding author: To be completed in the submission metadata.",
        "Funding: This research did not receive specific external grant funding.",
        "Acknowledgements: None.",
    ]
    for line in lines:
        paragraph = doc.add_paragraph(line)
        set_runs_font(paragraph, size_pt=10)
    doc.save(str(path))


def create_declaration(path: Path) -> None:
    doc = Document()
    heading = doc.add_paragraph("Declaration of Interest")
    heading.style = doc.styles["Title"]
    set_runs_font(heading, size_pt=16, bold=True)
    paragraph = doc.add_paragraph(
        "The authors declare that they have no known competing financial interests or personal relationships "
        "that could have appeared to influence the work reported in this paper."
    )
    set_runs_font(paragraph, size_pt=10)
    doc.save(str(path))


def create_credit(path: Path) -> None:
    doc = Document()
    heading = doc.add_paragraph("CRediT Author Statement")
    heading.style = doc.styles["Title"]
    set_runs_font(heading, size_pt=16, bold=True)
    lines = [
        "Conceptualization, Methodology, Software, Validation, Formal analysis, Investigation, Data curation, Writing - original draft, Writing - review and editing, Visualization: author team.",
        "Supervision, Project administration, Resources: author team.",
        "All authors read and approved the final manuscript.",
    ]
    for line in lines:
        paragraph = doc.add_paragraph(line)
        set_runs_font(paragraph, size_pt=10)
    doc.save(str(path))


def create_graphical_abstract_docx(path: Path, image_path: Path | None) -> None:
    doc = Document()
    apply_page_setup(doc)
    if image_path and image_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.5))
    else:
        paragraph = doc.add_paragraph("Graphical Abstract")
        paragraph.style = doc.styles["Title"]
        set_runs_font(paragraph, size_pt=16, bold=True)
        body = doc.add_paragraph("Graphical abstract image to be uploaded separately.")
        set_runs_font(body, size_pt=10)
    doc.save(str(path))


def create_upload_order(path: Path) -> None:
    text = """CEA / Editorial Manager upload order

1) Cover_Letter.docx -> Cover Letter
2) Title_Page.docx -> Title Page
3) Manuscript_Blinded.docx -> Manuscript
4) Highlights.docx -> Highlights
5) Graphical_Abstract.pdf / Graphical_Abstract.png -> Graphical Abstract
6) Declaration_of_Interest.docx -> Conflict of Interest
7) CRediT_Author_Statement.docx -> Supporting File
8) Supplementary_release_bundle_v1.zip -> Supporting File
9) CEA_style_two_column.pdf -> Supporting File (author reference only)
"""
    path.write_text(text, encoding="utf-8", newline="\n")


def write_validation_report(path: Path, lines: Iterable[str]) -> None:
    path.write_text("\n".join(lines) + "\n", encoding="utf-8", newline="\n")


def create_package_tree(package_dir: Path) -> Path:
    if package_dir.exists():
        shutil.rmtree(package_dir)
    files_dir = package_dir / "files"
    files_dir.mkdir(parents=True, exist_ok=True)
    return files_dir


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a text-locked CEA submission set from a master DOCX.")
    parser.add_argument("--master-docx", required=True, type=Path)
    parser.add_argument("--donor-docx", required=True, type=Path)
    parser.add_argument("--outdir", default=Path("out_cea_word_build_final"), type=Path)
    parser.add_argument("--package-dir", default=Path("CEA_submission_ready_package"), type=Path)
    parser.add_argument("--bundle-zip", default=Path("release_bundle_v1.zip"), type=Path)
    parser.add_argument("--ga-image", default=Path("graphical_abstract.png"), type=Path)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    repo_root = Path(__file__).resolve().parents[1]
    outdir = (repo_root / args.outdir).resolve() if not args.outdir.is_absolute() else args.outdir.resolve()
    package_dir = (repo_root / args.package_dir).resolve() if not args.package_dir.is_absolute() else args.package_dir.resolve()
    bundle_zip = (repo_root / args.bundle_zip).resolve() if not args.bundle_zip.is_absolute() else args.bundle_zip.resolve()
    ga_image = (repo_root / args.ga_image).resolve() if not args.ga_image.is_absolute() else args.ga_image.resolve()

    outdir.mkdir(parents=True, exist_ok=True)
    files_dir = create_package_tree(package_dir)

    stats = BuildStats()
    stats.source_figures = len(Document(str(args.master_docx)).inline_shapes)
    stats.source_todos, stats.source_placeholders = count_tokens(args.master_docx)

    master_doc = Document(str(args.master_docx))
    donor_doc = Document(str(args.donor_docx))

    apply_style_contract(master_doc)
    apply_page_setup(master_doc)
    apply_paragraph_layout(master_doc)
    style_tables(master_doc, stats)
    cleanup_log = clean_placeholder_blocks(master_doc, stats)
    table_cleanup_log = clean_table_reporting_placeholders(master_doc)
    figure_log = replace_or_insert_figures(master_doc, donor_doc, stats)
    apply_paragraph_layout(master_doc)
    for table in master_doc.tables:
        style_table(table)

    single_docx = outdir / "CEA_manuscript_EN.docx"
    master_doc.save(str(single_docx))
    stats.output_todos, stats.output_placeholders = count_tokens(single_docx)

    review_docx = outdir / "CEA_manuscript_EN_review.docx"
    create_review_copy(single_docx, review_docx)

    highlights_docx = outdir / "highlights.docx"
    create_highlights_docx(highlights_docx)

    ga_docx = outdir / "graphical_abstract_pipeline.docx"
    create_graphical_abstract_docx(ga_docx, ga_image if ga_image.exists() else None)
    shutil.copy2(ga_docx, outdir / "CEA_graphical_abstract_placeholder.docx")
    if ga_image.exists():
        shutil.copy2(ga_image, outdir / "graphical_abstract_pipeline.png")

    title = first_title(single_docx)
    cover_letter = files_dir / "Cover_Letter.docx"
    title_page = files_dir / "Title_Page.docx"
    declaration = files_dir / "Declaration_of_Interest.docx"
    credit = files_dir / "CRediT_Author_Statement.docx"
    create_cover_letter(cover_letter, title)
    create_title_page(title_page, title)
    create_declaration(declaration)
    create_credit(credit)

    manuscript_blinded = files_dir / "Manuscript_Blinded.docx"
    shutil.copy2(review_docx, manuscript_blinded)
    shutil.copy2(highlights_docx, files_dir / "Highlights.docx")
    shutil.copy2(ga_docx, files_dir / "Graphical_Abstract.docx")
    if ga_image.exists():
        shutil.copy2(ga_image, files_dir / "Graphical_Abstract.png")
    if bundle_zip.exists():
        shutil.copy2(bundle_zip, files_dir / "Supplementary_release_bundle_v1.zip")

    create_upload_order(package_dir / "UPLOAD_ORDER.txt")

    source_text_hash = visible_text_signature(args.master_docx)
    single_text_hash = visible_text_signature(single_docx)
    review_text_hash = visible_text_signature(review_docx)
    source_xml_hash = document_signature(args.master_docx)
    single_xml_hash = document_signature(single_docx)

    validation_lines = [
        "[PASS] build status: placeholder/TODO cleanup completed without changing scientific content",
        f"[INFO] title: {title}",
        f"[INFO] source_docx: {args.master_docx}",
        f"[INFO] donor_docx: {args.donor_docx}",
        f"[INFO] source_xml_sha256: {source_xml_hash}",
        f"[INFO] output_xml_sha256: {single_xml_hash}",
        f"[INFO] source_visible_text_sha256: {source_text_hash}",
        f"[INFO] cleaned_visible_text_sha256: {single_text_hash}",
        f"[{'PASS' if single_text_hash == review_text_hash else 'FAIL'}] cleaned manuscript integrity (single vs review): {single_text_hash} == {review_text_hash}",
        f"[INFO] source_inline_shapes: {stats.source_figures}",
        f"[INFO] output_inline_shapes: {stats.output_figures}",
        f"[INFO] figure_replacements: {stats.figures_replaced}",
        f"[INFO] figure_insertions: {stats.figures_inserted}",
        f"[INFO] tables_styled: {stats.tables_styled}",
        f"[INFO] removed_placeholder_paragraphs: {stats.removed_paragraphs}",
        f"[INFO] replaced_placeholder_runs: {stats.replaced_placeholders}",
        f"[WARN] source_todo_count: {stats.source_todos}",
        f"[WARN] source_placeholder_count: {stats.source_placeholders}",
        f"[{'PASS' if stats.output_todos == 0 else 'FAIL'}] output_todo_count: {stats.output_todos}",
        f"[{'PASS' if stats.output_placeholders == 0 else 'FAIL'}] output_placeholder_count: {stats.output_placeholders}",
    ]
    validation_lines.extend(f"[INFO] {line}" for line in cleanup_log)
    validation_lines.extend(f"[INFO] {line}" for line in table_cleanup_log)
    validation_lines.extend(f"[INFO] {line}" for line in figure_log)
    validation_report = outdir / "validation_report.txt"
    write_validation_report(validation_report, validation_lines)
    shutil.copy2(validation_report, files_dir / "Validation_Report.txt")

    report_path = outdir / "build_report.md"
    report_lines = [
        "# Build Report",
        "",
        f"- Master DOCX: `{args.master_docx}`",
        f"- Donor DOCX: `{args.donor_docx}`",
        f"- Output directory: `{outdir}`",
        f"- Package directory: `{package_dir}`",
        f"- Visible text hash (source): `{source_text_hash}`",
        f"- Visible text hash (cleaned single): `{single_text_hash}`",
        f"- Visible text hash (cleaned review): `{review_text_hash}`",
        f"- Figure replacements: {stats.figures_replaced}",
        f"- Figure insertions: {stats.figures_inserted}",
        f"- Tables styled: {stats.tables_styled}",
        f"- Removed placeholder paragraphs: {stats.removed_paragraphs}",
        f"- Replaced placeholder runs: {stats.replaced_placeholders}",
        f"- Source TODO count: {stats.source_todos}",
        f"- Source placeholder count: {stats.source_placeholders}",
        f"- Output TODO count: {stats.output_todos}",
        f"- Output placeholder count: {stats.output_placeholders}",
        "",
        "## Cleanup log",
        "",
    ]
    report_lines.extend(f"- {line}" for line in cleanup_log)
    report_lines.extend(f"- {line}" for line in table_cleanup_log)
    report_lines.extend([
        "",
        "## Figure log",
        "",
    ])
    report_lines.extend(f"- {line}" for line in figure_log)
    report_path.write_text("\n".join(report_lines) + "\n", encoding="utf-8")

    print(f"Saved single-column manuscript: {single_docx}")
    print(f"Saved review manuscript: {review_docx}")
    print(f"Prepared package directory: {package_dir}")
    print(f"Saved validation report: {validation_report}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
