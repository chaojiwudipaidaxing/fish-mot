from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


HEADING_TRANSLATIONS = {
    "Abstract": "摘要（中文工作注释）",
    "Introduction": "引言（中文工作注释）",
    "Related work": "相关工作（中文工作注释）",
    "Methods: investigator-developed evaluation framework": "方法：研究者主导评估框架（中文工作注释）",
    "Experimental setup": "实验设置（中文工作注释）",
    "Results": "结果（中文工作注释）",
    "Discussion: deployment guideline": "讨论：部署建议（中文工作注释）",
    "Cross-domain generalization protocol for minimal-cost validation": "跨域泛化分阶段方案（中文工作注释）",
    "Nomenclature / notation": "符号说明（中文工作注释）",
    "Limitations": "局限性（中文工作注释）",
    "Conclusion and future work": "结论与未来工作（中文工作注释）",
    "Data availability and reproducibility statement": "数据可用性与可复现性声明（中文工作注释）",
    "References": "参考文献（中文工作注释）",
}


def insert_paragraph_after(paragraph: Paragraph, text: str, style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name and style_name in paragraph.part.document.styles:
        new_para.style = paragraph.part.document.styles[style_name]
    run = new_para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    return new_para


def should_annotate_body(para: Paragraph) -> bool:
    if not para.text.strip():
        return False
    if not para.style:
        return False
    if para.style.name in {"Bibliography", "Image Caption", "Table Caption", "Captioned Figure"}:
        return False
    if para.style.name.startswith("Heading"):
        return False
    if len(para.text.strip()) < 30:
        return False
    return True


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create Chinese working DOCX with inline annotations.")
    parser.add_argument("--input", default="CEA_manuscript_EN.docx", help="Input English DOCX.")
    parser.add_argument(
        "--output",
        default="CEA_manuscript_ZH_working.docx",
        help="Output Chinese working DOCX.",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    in_path = Path(args.input)
    out_path = Path(args.output)
    doc = Document(in_path)

    heading_targets: list[Paragraph] = []
    body_targets: list[Paragraph] = []
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            heading_targets.append(para)
        elif should_annotate_body(para):
            body_targets.append(para)

    for para in reversed(heading_targets):
        heading_text = para.text.strip()
        zh = HEADING_TRANSLATIONS.get(
            heading_text,
            f"【中文标题注释】待翻译并润色本节标题：{heading_text}",
        )
        insert_paragraph_after(para, zh, "Body Text")

    for para in reversed(body_targets):
        insert_paragraph_after(
            para,
            "【中文工作注释】请将上段翻译为中文并校对术语一致性（drift-aware monitoring / "
            "audit-ready reproducibility / Scope A / Scope B / gating activation diagnostics）。",
            "Body Text",
        )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"saved={out_path}")
    print(f"headings_annotated={len(heading_targets)}")
    print(f"body_annotated={len(body_targets)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
