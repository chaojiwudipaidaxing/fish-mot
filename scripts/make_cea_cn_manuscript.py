from __future__ import annotations

import argparse
import re
import time
from pathlib import Path
from typing import Dict, List, Tuple

from deep_translator import GoogleTranslator
from docx import Document
from docx.text.paragraph import Paragraph


HEADING_MAP = {
    "Abstract": "\u6458\u8981",
    "Keywords": "\u5173\u952e\u8bcd",
    "Introduction": "\u5f15\u8a00",
    "Related work": "\u76f8\u5173\u5de5\u4f5c",
    "Tracking baselines and association design": "\u8ddf\u8e2a\u57fa\u7ebf\u4e0e\u5173\u8054\u8bbe\u8ba1",
    "Evaluation metrics and reproducible benchmarking": "\u8bc4\u4f30\u6307\u6807\u4e0e\u53ef\u590d\u73b0\u57fa\u51c6",
    "Aquaculture vision, fish tracking, and counting": "\u6c34\u4ea7\u89c6\u89c9\uff1a\u9c7c\u7c7b\u8ddf\u8e2a\u4e0e\u8ba1\u6570",
    "Methods: investigator-developed evaluation framework": "\u65b9\u6cd5\uff1a\u7814\u7a76\u8005\u4e3b\u5bfc\u7684\u8bc4\u6d4b\u6846\u67b6",
    "Terminology and scope": "\u672f\u8bed\u4e0e\u8303\u56f4",
    "Audit-ready reproducibility core": "\u5ba1\u8ba1\u5c31\u7eea\u7684\u53ef\u590d\u73b0\u6027\u6838\u5fc3",
    "Drift definition and closed-loop handling": "\u6f02\u79fb\u5b9a\u4e49\u4e0e\u95ed\u73af\u5904\u7406",
    "Input-layer indicator.": "\u8f93\u5165\u5c42\u6307\u6807",
    "Output-layer indicator.": "\u8f93\u51fa\u5c42\u6307\u6807",
    "Trigger and calibration.": "\u89e6\u53d1\u4e0e\u6821\u51c6",
    "Mitigation policy.": "\u7f13\u89e3\u7b56\u7565",
    "Controlled degradation protocol": "\u53ef\u63a7\u964d\u8d28\u534f\u8bae",
    "Environment-aware stress tests (underwater/aquaculture)": "\u73af\u5883\u611f\u77e5\u538b\u529b\u6d4b\u8bd5\uff08\u6c34\u4e0b/\u6c34\u4ea7\uff09",
    "Execution protocol.": "\u6267\u884c\u534f\u8bae",
    "Stratified difficulty diagnostics": "\u5206\u5c42\u96be\u5ea6\u8bca\u65ad",
    "Count stability and runtime profiling": "\u8ba1\u6570\u7a33\u5b9a\u6027\u4e0e\u8fd0\u884c\u65f6\u5256\u6790",
    "Framework visualization map for reviewers": "\u5ba1\u7a3f\u53cb\u597d\u7684\u6846\u67b6\u53ef\u89c6\u5316\u56fe\u8c31",
    "Experimental setup": "\u5b9e\u9a8c\u8bbe\u7f6e",
    "Frozen configuration and data": "\u51bb\u7ed3\u914d\u7f6e\u4e0e\u6570\u636e",
    "Methods and fairness protocol": "\u65b9\u6cd5\u4e0e\u516c\u5e73\u6027\u534f\u8bae",
    "Runtime measurement protocol": "\u8fd0\u884c\u65f6\u6d4b\u91cf\u534f\u8bae",
    "Research questions": "\u7814\u7a76\u95ee\u9898",
    "Results": "\u7ed3\u679c",
    "Core method family (Base to +adaptive)": "\u6838\u5fc3\u65b9\u6cd5\u65cf\uff08Base \u5230 +adaptive\uff09",
    "Strong baselines and combined comparison": "\u5f3a\u57fa\u7ebf\u4e0e\u7ec4\u5408\u5bf9\u6bd4",
    "Robustness under controlled degradation": "\u53ef\u63a7\u964d\u8d28\u4e0b\u7684\u9c81\u68d2\u6027",
    "Environment-aware stress tests (fixed 33 matrix)": "\u73af\u5883\u611f\u77e5\u538b\u529b\u6d4b\u8bd5\uff08\u56fa\u5b9a 3x3 \u77e9\u9635\uff09",
    "Gating threshold sensitivity (deployment robustness check)": "\u95e8\u63a7\u9608\u503c\u654f\u611f\u6027\uff08\u90e8\u7f72\u9c81\u68d2\u6027\u68c0\u67e5\uff09",
    "Gating activation analysis": "\u95e8\u63a7\u6fc0\u6d3b\u8bca\u65ad",
    "Drift-loop closed-loop validation (E1/E2/E3)": "\u6f02\u79fb\u95ed\u73af\u9a8c\u8bc1\uff08E1/E2/E3\uff09",
    "Stratified difficulty and failure modes": "\u5206\u5c42\u96be\u5ea6\u4e0e\u5931\u6548\u6a21\u5f0f",
    "Counting stability": "\u8ba1\u6570\u7a33\u5b9a\u6027",
    "Runtime and resource profile": "\u8fd0\u884c\u65f6\u4e0e\u8d44\u6e90\u5256\u6790",
    "Stage-time breakdown and deployment implication.": "\u9636\u6bb5\u8017\u65f6\u62c6\u89e3\u4e0e\u90e8\u7f72\u542b\u4e49",
    "Discussion: deployment guideline": "\u8ba8\u8bba\uff1a\u90e8\u7f72\u5efa\u8bae",
    "Cross-domain generalization protocol for minimal-cost validation": "\u8de8\u57df\u6cdb\u5316\u4e0e\u4f4e\u6210\u672c\u9a8c\u8bc1\u65b9\u6848",
    "Reporting items.": "\u62a5\u544a\u6761\u76ee",
    "Failure interpretation protocol.": "\u5931\u6548\u89e3\u91ca\u534f\u8bae",
    "Nomenclature / notation": "\u7b26\u53f7\u8bf4\u660e",
    "Limitations": "\u5c40\u9650\u6027",
    "Conclusion and future work": "\u7ed3\u8bba\u4e0e\u672a\u6765\u5de5\u4f5c",
    "Data availability and reproducibility statement": "\u6570\u636e\u53ef\u7528\u6027\u4e0e\u53ef\u590d\u73b0\u6027\u58f0\u660e",
    "Appendix A. Release package": "\u9644\u5f55 A\uff1a\u53d1\u5e03\u5305",
    "Funding": "\u57fa\u91d1\u8d44\u52a9",
    "Author contributions (CRediT)": "\u4f5c\u8005\u8d21\u732e\uff08CRediT\uff09",
    "Competing interests": "\u5229\u76ca\u51b2\u7a81",
    "Declaration of Generative AI and AI-assisted technologies in the manuscript preparation process": "\u7a3f\u4ef6\u64b0\u5199\u8fc7\u7a0b\u4e2d\u751f\u6210\u5f0f AI \u4e0e AI \u8f85\u52a9\u6280\u672f\u58f0\u660e",
}

TOKEN_PATTERNS = [
    re.compile(r"https?://[^\s]+"),
    re.compile(r"\b[\w./\\-]+\.(?:csv|tex|pdf|png|jpg|jpeg|docx|json|txt|md|py|sh)\b", re.IGNORECASE),
    re.compile(r"\b[A-Za-z]+(?:_[A-Za-z0-9]+)+(?:\([^)]+\))?\b"),
    re.compile(r"\b[A-Za-z]+\([^)]+\)"),
    re.compile(r"\\[A-Za-z]+(?:\{[^}]*\})*"),
    re.compile(r"\b(?:HOTA|IDF1|IDSW|DetA|AssA|CountMAE|CountRMSE|CountVar|CountDrift|FPS|CPU|RSS)\b"),
    re.compile(r"\[[0-9,\-\s]+\]"),
]

CODE_ONLY_PAT = re.compile(r"^[\sA-Za-z0-9_./\\\-+*=<>()[\]{}:,;|`~^%$#@!?&]+$")
HAS_LATIN_PAT = re.compile(r"[A-Za-z]")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate CEA-style Chinese manuscript DOCX from English DOCX.")
    parser.add_argument("--input", default="out_cea_word_build_final/CEA_manuscript_EN.docx")
    parser.add_argument("--output", default="out_cea_word_build_cn/CEA_manuscript_CN_cea_style.docx")
    parser.add_argument("--max-retries", type=int, default=3)
    parser.add_argument("--sleep-sec", type=float, default=0.25)
    return parser.parse_args()


def is_references_heading(text: str) -> bool:
    t = text.strip().lower()
    return t in {"references", "reference"}


def is_formula_paragraph(paragraph: Paragraph) -> bool:
    xml = paragraph._p.xml  # noqa: SLF001
    return ("<m:oMath" in xml) or ("<m:oMathPara" in xml)


def protect_tokens(text: str) -> Tuple[str, List[str]]:
    tokens: List[str] = []

    def repl(match: re.Match[str]) -> str:
        tokens.append(match.group(0))
        return f"__PH_{len(tokens)-1}__"

    out = text
    for pat in TOKEN_PATTERNS:
        out = pat.sub(repl, out)
    return out, tokens


def restore_tokens(text: str, tokens: List[str]) -> str:
    out = text
    for i, tok in enumerate(tokens):
        out = out.replace(f"__PH_{i}__", tok)
    return out


def post_polish_zh(text: str) -> str:
    replacements = {
        "drift-proof": "drift-aware monitoring",
        "Drift-proof": "drift-aware monitoring",
    }
    out = text
    for k, v in replacements.items():
        out = out.replace(k, v)
    return out


def should_skip_run(run_text: str) -> bool:
    t = run_text.strip()
    if not t:
        return True
    if not HAS_LATIN_PAT.search(t):
        return True
    if CODE_ONLY_PAT.match(t) and len(t.split()) <= 4:
        return True
    return False


def translate_text(
    text: str,
    translator: GoogleTranslator,
    cache: Dict[str, str],
    max_retries: int,
    sleep_sec: float,
) -> str:
    if text in cache:
        return cache[text]
    protected, tokens = protect_tokens(text)
    if not HAS_LATIN_PAT.search(protected):
        return text

    last_exc: Exception | None = None
    for _ in range(max_retries):
        try:
            zh = translator.translate(protected)
            zh = restore_tokens(zh, tokens)
            zh = post_polish_zh(zh)
            cache[text] = zh
            return zh
        except Exception as exc:
            last_exc = exc
            time.sleep(sleep_sec)

    if last_exc is not None:
        print(f"[WARN] translate failed, keep original: {text[:90]!r} :: {last_exc}")
    cache[text] = text
    return text


def apply_heading_map(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if text in HEADING_MAP:
        mapped = HEADING_MAP[text]
        if paragraph.runs:
            paragraph.runs[0].text = mapped
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(mapped)
        return True

    if text == "Anonymous Author(s)":
        mapped = "\u533f\u540d\u4f5c\u8005"
        if paragraph.runs:
            paragraph.runs[0].text = mapped
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(mapped)
        return True

    return False


def main() -> int:
    args = parse_args()
    in_path = Path(args.input)
    out_path = Path(args.output)
    if not in_path.exists():
        raise FileNotFoundError(f"input docx not found: {in_path}")

    doc = Document(str(in_path))
    translator = GoogleTranslator(source="en", target="zh-CN")
    cache: Dict[str, str] = {}

    in_references = False
    changed_runs = 0
    changed_headings = 0
    skipped_formula_paragraphs = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""

        if is_references_heading(text):
            in_references = True
            continue
        if in_references or style_name == "Bibliography":
            continue

        if apply_heading_map(para):
            changed_headings += 1
            continue

        if is_formula_paragraph(para):
            skipped_formula_paragraphs += 1
            continue

        for run in para.runs:
            raw = run.text
            if should_skip_run(raw):
                continue
            zh = translate_text(raw, translator, cache, args.max_retries, args.sleep_sec)
            if zh != raw:
                run.text = zh
                changed_runs += 1

    # second pass: paragraph-level cleanup for mixed-language leftovers
    in_references = False
    cleaned_paragraphs = 0
    paragraph_cache: Dict[str, str] = {}

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""

        if is_references_heading(text):
            in_references = True
            continue
        if in_references or style_name == "Bibliography":
            continue
        if not text or not HAS_LATIN_PAT.search(text):
            continue
        if (not style_name.startswith("Heading")) and CODE_ONLY_PAT.match(text) and len(text.split()) <= 6:
            continue

        if text in paragraph_cache:
            zh = paragraph_cache[text]
        else:
            zh = translate_text(text, translator, cache, args.max_retries, args.sleep_sec)
            paragraph_cache[text] = zh

        if zh != text:
            if para.runs:
                para.runs[0].text = zh
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.add_run(zh)
            cleaned_paragraphs += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"saved={out_path}")
    print(f"changed_runs={changed_runs}")
    print(f"changed_headings={changed_headings}")
    print(f"skipped_formula_paragraphs={skipped_formula_paragraphs}")
    print(f"cleaned_paragraphs={cleaned_paragraphs}")
    print(f"cache_size={len(cache)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
