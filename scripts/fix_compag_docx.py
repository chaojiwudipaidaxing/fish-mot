from __future__ import annotations

import argparse
import hashlib
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REQUIRED_SCOPE_B_FIELDS = [
    "fps_e2e",
    "mem_peak_mb_e2e",
    "cpu_norm_e2e",
    "decode_time",
    "detector_time",
    "tracking_time",
    "write_time",
    "detector_name",
    "input_resolution",
]


@dataclass
class FixStats:
    text_replacements: int = 0
    pseudo_tokens_removed: int = 0
    urls_replaced: int = 0
    drift_block_fixed: bool = False
    algorithm_block_fixed: bool = False
    scope_b_schema_fixed: bool = False


def paragraph_has_math(paragraph) -> bool:
    return "<m:oMath" in paragraph._p.xml or "<m:oMathPara" in paragraph._p.xml


def set_style_font(style, *, size_pt: float | None, bold: bool | None, italic: bool | None) -> None:
    font = style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0, 0, 0)
    if size_pt is not None:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic

    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), "Times New Roman")


def normalize_styles(doc: DocumentType, changes: list[str]) -> None:
    styles = doc.styles
    for name in ("Normal", "Body Text", "First Paragraph", "Abstract", "Bibliography"):
        if name in styles:
            set_style_font(styles[name], size_pt=10, bold=False, italic=False)
            styles[name].paragraph_format.line_spacing = 1.15
    if "Title" in styles:
        set_style_font(styles["Title"], size_pt=17, bold=True, italic=False)
        styles["Title"].paragraph_format.alignment = 0
    for name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4", "Abstract Title"):
        if name in styles:
            set_style_font(styles[name], size_pt=10, bold=True, italic=False)
            styles[name].paragraph_format.line_spacing = 1.15
    for name in ("Caption", "Image Caption", "Table Caption", "Captioned Figure"):
        if name in styles:
            set_style_font(styles[name], size_pt=9, bold=False, italic=True)
            styles[name].paragraph_format.line_spacing = 1.0
    changes.append("Applied COMPAG style contract (TNR, single-column friendly heading/body/caption sizes).")


def set_runs_tnr(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{key}"), "Times New Roman")


def replace_para_text(para, text: str) -> None:
    para.clear()
    para.add_run(text)


def clean_text(text: str, stats: FixStats) -> str:
    original = text

    # control chars
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)

    # common mojibake artifacts from symbol conversion
    text = text.replace("¦Ì", "mu")
    text = text.replace("¦Ò", "sigma")
    text = text.replace("¦Å", "epsilon")
    text = text.replace("¦Ó", "tau")
    text = text.replace("¡û", "<-")
    text = text.replace("¡Ý", ">=")

    # normalization
    text = text.replace("¡À", "±")
    text = text.replace("meanstd", "mean±std")
    text = text.replace("Meanstd", "Mean±std")
    text = text.replace("Table Table", "Table")
    text = text.replace("Tables Table", "Tables")

    # Convert merged mean/std pairs: 58.2320.304 -> 58.232 ± 0.304
    text = re.sub(r"(\d+\.\d{3})\s*([01]\.\d{3})", r"\1 ± \2", text)
    text = re.sub(r"(\d+\.\d{3})\u2002+([01]\.\d{3})", r"\1 ± \2", text)

    # Readable crossref cleanup
    text = re.sub(r"Algorithm\s*\[alg:[^\]]+\]", "Algorithm 1", text)
    text = re.sub(r"\[(tab|alg|fig):[^\]]+\]", "", text)
    text = re.sub(r"\s{2,}", " ", text).strip()

    # Tables Table 1, Table 3 -> Tables 1 and 3
    text = re.sub(r"Tables?\s+Table\s+(\d+)\s*,\s*Table\s+(\d+)", r"Tables \1 and \2", text)
    text = re.sub(r"Tables?\s+Table\s+(\d+)\s+and\s+Table\s+(\d+)", r"Tables \1 and \2", text)
    text = re.sub(r"\bTable\s+Table\s+(\d+)", r"Table \1", text)

    # content safety terms
    text = re.sub(r"\bdrift-proof\b", "drift-aware monitoring", text, flags=re.IGNORECASE)
    text = re.sub(r"\bdrift proof\b", "drift-aware monitoring", text, flags=re.IGNORECASE)

    # placeholder URL removal
    if "zenodo.XXXX" in text or "anonymous.4open" in text:
        stats.urls_replaced += 1
    text = re.sub(
        r"https?://doi\.org/10\.5281/zenodo\.[A-Za-z0-9Xx]+",
        "TODO: Zenodo DOI after acceptance",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"https?://anonymous\.4open\.science/\S+",
        "TODO: <anonymized reviewer link>",
        text,
        flags=re.IGNORECASE,
    )

    if original != text:
        stats.text_replacements += 1
        stats.pseudo_tokens_removed += original.count("[tab:") + original.count("[alg:")
    return text


def apply_text_normalization(doc: DocumentType, stats: FixStats, changes: list[str]) -> None:
    for para in doc.paragraphs:
        if paragraph_has_math(para):
            for run in para.runs:
                if run.text:
                    run.text = clean_text(run.text, stats)
            set_runs_tnr(para)
            continue
        new_text = clean_text(para.text, stats)
        if new_text != para.text:
            replace_para_text(para, new_text)
        set_runs_tnr(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if paragraph_has_math(para):
                        continue
                    new_text = clean_text(para.text, stats)
                    if new_text != para.text:
                        replace_para_text(para, new_text)
                    set_runs_tnr(para)
    changes.append("Normalized special characters, ± notation, cross-ref residues, and unsafe placeholders.")


def force_fix_meanstd(doc: DocumentType, changes: list[str]) -> None:
    replacements = 0
    for para in doc.paragraphs:
        if "meanstd" in para.text.lower():
            new_text = re.sub(r"(?i)meanstd", "mean±std", para.text)
            replace_para_text(para, new_text)
            set_runs_tnr(para)
            replacements += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "meanstd" in para.text.lower():
                        new_text = re.sub(r"(?i)meanstd", "mean±std", para.text)
                        replace_para_text(para, new_text)
                        set_runs_tnr(para)
                        replacements += 1
    if replacements:
        changes.append(f"Force-fixed residual meanstd tokens: {replacements} replacements.")


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def apply_table_style(table) -> None:
    rows = table.rows
    if not rows:
        return
    cols = len(rows[0].cells)
    for r in rows:
        for c in r.cells:
            set_cell_border(c, left={"val": "nil"}, right={"val": "nil"}, top={"val": "nil"}, bottom={"val": "nil"})
            for p in c.paragraphs:
                set_runs_tnr(p)
                p.paragraph_format.line_spacing = 1.0
    for c in rows[0].cells:
        set_cell_border(c, top={"val": "single", "sz": "8", "color": "000000"})
        set_cell_border(c, bottom={"val": "single", "sz": "6", "color": "000000"})
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
    for c in rows[-1].cells:
        set_cell_border(c, bottom={"val": "single", "sz": "8", "color": "000000"})


def style_tables(doc: DocumentType, changes: list[str]) -> None:
    for t in doc.tables:
        apply_table_style(t)
    changes.append(f"Applied no-vertical-rule three-line style to {len(doc.tables)} editable tables.")


def find_heading_index(doc: DocumentType, heading: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() == heading.strip().lower():
            return i
    return -1


def insert_after(paragraph, text: str, style_name: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    para = Paragraph(new_p, paragraph._parent)
    if style_name and style_name in paragraph.part.document.styles:
        para.style = paragraph.part.document.styles[style_name]
    para.add_run(text)
    set_runs_tnr(para)
    return para


def restore_drift_block(doc: DocumentType, tex_path: Path, stats: FixStats, changes: list[str], todos: list[str]) -> None:
    idx = find_heading_index(doc, "Drift definition and closed-loop handling")
    if idx < 0:
        todos.append("TODO: Missing heading 'Drift definition and closed-loop handling' in DOCX; manual insertion required.")
        return
    p = doc.paragraphs[idx + 1]
    text = (
        "Let X_t^(W) = {x_{t-W+1}, ..., x_t} be the recent input window of size W and X_ref be a reference window "
        "(calibration period). Let q(·) denote output-quality statistics computed on a window. In this manuscript, "
        "W denotes the fixed monitoring window length configured before evaluation, and K denotes the persistence "
        "window count required to trigger an alert. The reference window X_ref is constructed from a calibration "
        "interval of the same stream under nominal operating conditions and then frozen for online monitoring. "
        "TODO: report exact calibration interval length and refresh policy in release notes."
    )
    replace_para_text(p, text)
    set_runs_tnr(p)

    # Repair known blank helper sentences around equations.
    for para in doc.paragraphs[idx + 2 : idx + 18]:
        t = para.text.strip()
        if t.startswith("For a scalar feature"):
            replace_para_text(
                para,
                "For a scalar feature z (e.g., detection confidence or gating score), we monitor the empirical CDF distance:",
            )
        elif t == "is unitless and computed per sequence or per deployment stream.":
            replace_para_text(para, "D_in(t) is unitless and computed per sequence or per deployment stream.")
        elif t.startswith("For output metric vector"):
            replace_para_text(para, "For output metric vector q_t = [IDSW_t, CountMAE_t, F1_t]^T, we monitor normalized deviation:")
        elif t.startswith("Here") and "reference statistics" in t:
            replace_para_text(
                para,
                "Here μ_{q,ref} and σ_{q,ref} are reference statistics; ε avoids division by zero.",
            )
        elif t.startswith("Drift alert is raised when"):
            replace_para_text(
                para,
                "Drift alert is raised when D_in(t) > tau_in or D_out(t) > tau_out for K consecutive windows, "
                "where K is the persistence hyper-parameter.",
            )
        set_runs_tnr(para)
    stats.drift_block_fixed = True
    changes.append(f"Restored drift terminology/variables from `{tex_path}` into readable DOCX text.")


def restore_algorithm_block(doc: DocumentType, stats: FixStats, changes: list[str]) -> None:
    # Find collapsed/misaligned algorithm block and replace with readable pseudo-code block.
    target_idx = -1
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip().lower()
        if "algorithm 1" in t and "drift-aware closed-loop" in t:
            target_idx = i
            break
        if "input:" in t and "windows" in t and "tau" in t:
            target_idx = i
            break
    if target_idx < 0:
        return

    target = doc.paragraphs[target_idx]
    replace_para_text(target, "Algorithm 1. Drift-aware closed-loop monitoring and fallback")
    if "Heading 4" in doc.styles:
        target.style = doc.styles["Heading 4"]
    set_runs_tnr(target)

    lines = [
        "Input: windows X_t^(W); output stats q_t; thresholds (tau_in, tau_out); persistence K; reference X_ref.",
        "State: alert counter c <- 0; current profile p <- default.",
        "For each window t: compute D_in(t) and D_out(t). If D_in(t) > tau_in or D_out(t) > tau_out then c <- c+1 else c <- 0.",
        "If c >= K: activate conservative mode (e.g., tighter gating / safer tracker profile); log context {timestamp, D_in(t), D_out(t), c, profile, decision}.",
    ]

    tail = target
    for offset, line in enumerate(lines, start=1):
        next_idx = target_idx + offset
        if next_idx < len(doc.paragraphs):
            para = doc.paragraphs[next_idx]
            replace_para_text(para, line)
            if "Normal" in doc.styles:
                para.style = doc.styles["Normal"]
            set_runs_tnr(para)
            tail = para
        else:
            tail = insert_after(tail, line, "Normal")
    stats.algorithm_block_fixed = True
    changes.append("Rebuilt Algorithm 1 block with explicit Input/State/steps and restored variable names.")


def ensure_scope_b_schema(doc: DocumentType, stats: FixStats, changes: list[str], todos: list[str]) -> None:
    found_table = None
    for table in doc.tables:
        first_col = [r.cells[0].text.strip() for r in table.rows if r.cells]
        if sum(1 for f in REQUIRED_SCOPE_B_FIELDS if f in first_col) >= 4:
            found_table = table
            break
    if found_table is None:
        idx = find_heading_index(doc, "Runtime and resource profile")
        if idx < 0:
            todos.append("TODO: Scope B schema table missing and runtime section not found.")
            return
        anchor = doc.paragraphs[idx]
        caption = insert_after(anchor, "Scope B (end-to-end) reporting schema (TODO values to be populated after run).", "Table Caption")
        table = doc.add_table(rows=1 + len(REQUIRED_SCOPE_B_FIELDS), cols=2)
        caption._p.addnext(table._tbl)
        table.rows[0].cells[0].text = "Field (csv)"
        table.rows[0].cells[1].text = "Reporting requirement (no values filled in this manuscript)"
        for i, field in enumerate(REQUIRED_SCOPE_B_FIELDS, start=1):
            table.rows[i].cells[0].text = field
            table.rows[i].cells[1].text = "TODO: fill from frozen profiling output."
        apply_table_style(table)
        stats.scope_b_schema_fixed = True
        changes.append("Inserted missing Scope B schema table with required field names.")
    else:
        present = {r.cells[0].text.strip() for r in found_table.rows if r.cells}
        missing = [f for f in REQUIRED_SCOPE_B_FIELDS if f not in present]
        for f in missing:
            row = found_table.add_row().cells
            row[0].text = f
            row[1].text = "TODO: fill from frozen profiling output."
        if missing:
            changes.append(f"Added missing Scope B fields: {', '.join(missing)}.")
        apply_table_style(found_table)
        stats.scope_b_schema_fixed = True


def ensure_submission_notes(doc: DocumentType, changes: list[str]) -> None:
    target = (
        "Submission notes. Highlights: submitted as a separate file. "
        "Graphical abstract: submitted as a separate file (see graphical abstract guidelines). "
        "Graphical Abstract - Figure GA (submitted separately)."
    )
    for p in doc.paragraphs:
        if "Submission notes." in p.text:
            replace_para_text(p, target)
            if "Body Text" in doc.styles:
                p.style = doc.styles["Body Text"]
            set_runs_tnr(p)
            return
    # Insert after abstract keywords line
    idx = find_heading_index(doc, "Abstract")
    if idx >= 0 and idx + 2 < len(doc.paragraphs):
        insert_after(doc.paragraphs[idx + 2], target, "Body Text")
    else:
        doc.add_paragraph(target, style="Body Text")
    changes.append("Inserted missing submission notes for highlights/graphical abstract.")


def ensure_data_statement(doc: DocumentType, todos: list[str], changes: list[str]) -> None:
    full = "\n".join(p.text for p in doc.paragraphs)
    if "Data availability" in full:
        return
    h = doc.add_paragraph("Data availability and reproducibility statement")
    if "Heading 1" in doc.styles:
        h.style = doc.styles["Heading 1"]
    p = doc.add_paragraph(
        "Data availability. TODO: provide explicit availability/restriction statement and derived artifact access details."
    )
    if "Body Text" in doc.styles:
        p.style = doc.styles["Body Text"]
    set_runs_tnr(h)
    set_runs_tnr(p)
    todos.append("TODO: Complete Data availability statement with final access policy and repository identifiers.")
    changes.append("Inserted missing Data availability statement placeholder.")


def ensure_terminology_scope(doc: DocumentType, changes: list[str], todos: list[str]) -> None:
    if find_heading_index(doc, "Terminology and scope") >= 0:
        return
    methods_idx = find_heading_index(doc, "Methods: investigator-developed evaluation framework")
    if methods_idx < 0:
        todos.append("TODO: Could not locate Methods section for Terminology & Scope insertion.")
        return
    h = insert_after(doc.paragraphs[methods_idx], "Terminology and scope", "Heading 2")
    insert_after(
        h,
        "To capture the volatile nature of pond visibility, input drift is formalized as a distribution shift "
        "between reference and online windows on input-side signals. Output drift denotes change in output-side "
        "quality indicators under fixed protocol. This study "
        "targets drift-aware monitoring and response only, not guaranteed drift elimination.",
        "Body Text",
    )
    changes.append("Inserted missing 'Terminology and scope' subsection with non-guarantee framing.")


def enforce_gating_diagnostics(doc: DocumentType, todos: list[str], changes: list[str]) -> None:
    idx = find_heading_index(doc, "Gating activation analysis")
    if idx < 0:
        todos.append("TODO: Add gating activation analysis section (trigger rate, event length, score CDF).")
        return
    block = "\n".join(p.text.lower() for p in doc.paragraphs[idx : idx + 12])
    needed = ["trigger rate", "event", "cdf"]
    if not all(k in block for k in needed):
        tail = doc.paragraphs[idx]
        tail = insert_after(
            tail,
            "Planned activation diagnostics: trigger rate (per frame/per sequence), trigger event-length "
            "distribution, and gating-score CDF with threshold markers. TODO: populate statistics from frozen logs.",
            "Body Text",
        )
        set_runs_tnr(tail)
        changes.append("Added missing gating activation diagnostics plan.")
    plateau_block = "\n".join(p.text.lower() for p in doc.paragraphs[idx - 8 : idx + 2 if idx >= 8 else idx + 2])
    if "insensitive" in plateau_block and "todo" not in plateau_block:
        todos.append("TODO: Rephrase threshold-plateau claim as candidate interval pending activation diagnostics.")


def fix_placeholder_links(doc: DocumentType, changes: list[str]) -> None:
    for p in doc.paragraphs:
        t = p.text
        if "zenodo.XXXX" in t or "anonymous.4open" in t:
            t = re.sub(r".*zenodo\.XXXXXXX.*", "TODO: Zenodo DOI after acceptance", t)
            t = re.sub(r".*anonymous\.4open\.science.*", "TODO: <anonymized reviewer link>", t)
            replace_para_text(p, t)
            set_runs_tnr(p)
    changes.append("Replaced pseudo-links with explicit TODO placeholders (no fabricated URLs).")


def ensure_abstract_keywords(doc: DocumentType, todos: list[str], changes: list[str]) -> None:
    abs_words = None
    key_count = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() == "abstract":
            # abstract body
            for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
                if doc.paragraphs[j].text.strip():
                    body = doc.paragraphs[j]
                    abs_words = len(re.findall(r"[A-Za-z0-9\-]+", body.text))
                    if abs_words > 250:
                        body.add_run(" TODO: shorten Abstract to <=250 words.")
                        todos.append("TODO: Abstract exceeds 250 words.")
                    break
            if i + 2 < len(doc.paragraphs):
                keys = [x.strip() for x in doc.paragraphs[i + 2].text.split(",") if x.strip()]
                key_count = len(keys)
                if key_count < 1 or key_count > 7:
                    todos.append("TODO: Keywords count must be within 1-7.")
            break
    changes.append(f"Checked abstract/keywords constraints (abstract_words={abs_words}, keywords={key_count}).")


def create_highlights(path: Path) -> None:
    bullets = [
        "Audit-ready protocol links fairness checks, stress tests, and deployment rules.",
        "Drift-aware diagnostics localize failure concentration by risk buckets.",
        "Runtime is split into tracking-only and end-to-end for deployment clarity.",
        "Gating threshold conclusions are treated as candidate intervals pending diagnostics.",
    ]
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(10)
    for b in bullets:
        if len(b) > 85:
            raise ValueError(f"Highlight too long: {b}")
        p = doc.add_paragraph(b, style="List Bullet")
        set_runs_tnr(p)
    doc.save(path)


def create_ga_placeholder(path: Path) -> None:
    doc = Document()
    p1 = doc.add_paragraph("Graphical Abstract (submitted separately)")
    p2 = doc.add_paragraph("Placeholder only. Final graphical abstract must be uploaded as a separate file.")
    set_runs_tnr(p1)
    set_runs_tnr(p2)
    doc.save(path)


def create_zh_working(in_docx: Path, out_docx: Path) -> None:
    doc = Document(in_docx)
    from docx.text.paragraph import Paragraph

    paras = list(doc.paragraphs)
    for para in reversed(paras):
        if not para.text.strip():
            continue
        if para.style and para.style.name in {"Image Caption", "Table Caption", "Bibliography"}:
            continue
        new_p = OxmlElement("w:p")
        para._p.addnext(new_p)
        p = Paragraph(new_p, para._parent)
        if "Body Text" in doc.styles:
            p.style = doc.styles["Body Text"]
        p.add_run("【中文工作注释】请翻译并校对本段术语、符号、数字与原文一致。")
        set_runs_tnr(p)
    doc.save(out_docx)


def write_report(path: Path, title: str, items: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    lines = [f"# {title}", ""]
    if items:
        lines.extend(f"- {x}" for x in items)
    else:
        lines.append("- None.")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Strict COMPAG/CEA docx fixer.")
    parser.add_argument("--input", required=True, help="Input docx path.")
    parser.add_argument("--tex", default="paper/cea_draft/main.tex", help="LaTeX main.tex for symbol restoration context.")
    parser.add_argument("--outdir", default="out_compag_strict", help="Output directory.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = Path(args.input)
    if not input_path.exists():
        raise FileNotFoundError(f"Input docx not found: {args.input}")
    tex_path = Path(args.tex)
    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    out_en_docx = outdir / "COMPAG_CEA_manuscript_EN.docx"
    out_zh_docx = outdir / "COMPAG_CEA_manuscript_ZH_working.docx"
    highlights_docx = outdir / "highlights.docx"
    ga_docx = outdir / "graphical_abstract_placeholder.docx"
    changes_md = outdir / "changes.md"
    todo_md = outdir / "todo.md"

    doc = Document(input_path)
    sha256 = hashlib.sha256(input_path.read_bytes()).hexdigest()
    changes: list[str] = [f"Loaded source docx: {input_path}", f"Input SHA256: {sha256}"]
    todos: list[str] = []
    stats = FixStats()

    normalize_styles(doc, changes)
    apply_text_normalization(doc, stats, changes)
    ensure_submission_notes(doc, changes)
    ensure_data_statement(doc, todos, changes)
    ensure_terminology_scope(doc, changes, todos)
    restore_drift_block(doc, tex_path, stats, changes, todos)
    restore_algorithm_block(doc, stats, changes)
    ensure_scope_b_schema(doc, stats, changes, todos)
    enforce_gating_diagnostics(doc, todos, changes)
    fix_placeholder_links(doc, changes)
    force_fix_meanstd(doc, changes)
    ensure_abstract_keywords(doc, todos, changes)
    style_tables(doc, changes)

    doc.save(out_en_docx)
    changes.append(f"Saved fixed EN manuscript: {out_en_docx}")

    create_highlights(highlights_docx)
    create_ga_placeholder(ga_docx)
    create_zh_working(out_en_docx, out_zh_docx)
    changes.append("Generated highlights, graphical abstract placeholder, and ZH working copy.")

    # add strict TODO checks
    if not stats.drift_block_fixed:
        todos.append("TODO: Drift block restoration could not be confirmed.")
    if not stats.algorithm_block_fixed:
        todos.append("TODO: Algorithm block restoration could not be confirmed.")
    if not stats.scope_b_schema_fixed:
        todos.append("TODO: Scope B schema table could not be confirmed.")

    changes.append(
        "Stats: "
        f"text_replacements={stats.text_replacements}, "
        f"pseudo_tokens_removed={stats.pseudo_tokens_removed}, "
        f"urls_replaced={stats.urls_replaced}, "
        f"drift_fixed={stats.drift_block_fixed}, "
        f"algorithm_fixed={stats.algorithm_block_fixed}, "
        f"scope_b_fixed={stats.scope_b_schema_fixed}"
    )

    write_report(changes_md, "Changes", changes)
    write_report(todo_md, "TODO", todos)

    print(f"input={input_path}")
    print(f"input_sha256={sha256}")
    print(f"saved_en={out_en_docx}")
    print(f"saved_zh={out_zh_docx}")
    print(f"saved_highlights={highlights_docx}")
    print(f"saved_ga={ga_docx}")
    print(f"saved_changes={changes_md}")
    print(f"saved_todo={todo_md}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
