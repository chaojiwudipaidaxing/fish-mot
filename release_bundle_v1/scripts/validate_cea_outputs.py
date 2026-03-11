from __future__ import annotations

import argparse
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}


def load_doc_xml(path: Path) -> ET.Element:
    with zipfile.ZipFile(path) as z:
        return ET.fromstring(z.read("word/document.xml"))


def check_fonts_no_calibri(docx_path: Path) -> tuple[bool, str]:
    root = load_doc_xml(docx_path)
    bad = 0
    for r_fonts in root.findall(".//w:rFonts", NS):
        vals = " ".join(dict(r_fonts.attrib).values())
        if any(x in vals for x in ["Calibri", "minorHAnsi", "majorHAnsi"]):
            bad += 1
    return bad == 0, f"bad_font_refs={bad}"


def check_single_column(docx_path: Path) -> tuple[bool, str]:
    root = load_doc_xml(docx_path)
    bad = 0
    for cols in root.findall(".//w:sectPr/w:cols", NS):
        num = cols.attrib.get(f"{{{NS['w']}}}num")
        if num and num != "1":
            bad += 1
    return bad == 0, f"multi_column_sections={bad}"


def check_formula_integrity(docx_path: Path) -> tuple[bool, str]:
    root = load_doc_xml(docx_path)
    omml_count = len(root.findall(".//m:oMath", NS)) + len(root.findall(".//m:oMathPara", NS))
    let_blank = 0
    gated_blank = 0
    for p in root.findall(".//w:p", NS):
        text = "".join((t.text or "") for t in p.findall(".//w:t", NS))
        if re.search(r"Let\s{2,}be", text):
            let_blank += 1
        if "are masked when ." in text:
            gated_blank += 1
    ok = omml_count > 0 and let_blank == 0 and gated_blank == 0
    return ok, f"omml={omml_count} let_blank={let_blank} gated_blank={gated_blank}"


def check_abstract_keywords(docx_path: Path) -> tuple[bool, str]:
    doc = Document(docx_path)
    paras = [p.text.strip() for p in doc.paragraphs]
    abstract = ""
    keywords = ""
    for i, t in enumerate(paras):
        if t.lower() == "abstract" and i + 1 < len(paras):
            abstract = paras[i + 1].strip()
            if i + 2 < len(paras):
                keywords = paras[i + 2].strip()
            break
    abs_words = len(re.findall(r"[A-Za-z0-9\-]+", abstract))
    keys = [k.strip() for k in re.split(r",|;|\\sep", keywords) if k.strip()]
    ok = abs_words <= 250 and 1 <= len(keys) <= 7
    return ok, f"abstract_words={abs_words} keyword_count={len(keys)}"


def check_submission_notes(docx_path: Path) -> tuple[bool, str]:
    doc = Document(docx_path)
    full = "\n".join(p.text for p in doc.paragraphs)
    checks = {
        "highlights_note": "Highlights: submitted as a separate file." in full,
        "ga_note": "Graphical abstract: submitted as a separate file" in full,
        "ga_callout": ("Graphical Abstract" in full and "submitted separately" in full),
        "data_statement": "Data availability." in full or "Data availability" in full,
    }
    ok = all(checks.values())
    return ok, " ".join(f"{k}={v}" for k, v in checks.items())


def check_text_artifacts(docx_path: Path) -> tuple[bool, str]:
    doc = Document(docx_path)
    full = "\n".join(p.text for p in doc.paragraphs)
    patterns = {
        "table_table": r"Table Table",
        "tab_token": r"\[tab:",
        "alg_token": r"\[alg:",
        "meanstd": r"\bmeanstd\b",
        "let_be": r"Let\s{2,}be",
        "gated_blank": r"are masked when\s*\.",
        "todo": r"\bTODO\b",
        "placeholder": r"placeholder",
    }
    hits = {k: len(re.findall(v, full)) for k, v in patterns.items()}
    # Catch dropped ± where two decimals are separated only by spaces/thin spaces.
    pm_missing = len(re.findall(r"\d+\.\d{3}[\u2000-\u200B\u00A0 ]{1,}\d+\.\d{3}", full))
    hits["pm_missing"] = pm_missing
    ok = all(v == 0 for v in hits.values())
    return ok, " ".join(f"{k}={v}" for k, v in hits.items())


def check_tables_and_captions(docx_path: Path) -> tuple[bool, str]:
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    pseudo_hits = sum(text.count(pat) for pat in ["l S S", "Method & HOTA", " & " * 4])
    img_count = len(doc.inline_shapes)
    cap_count = sum(
        1
        for p in doc.paragraphs
        if p.style and p.style.name in {"Image Caption", "Caption"} and p.text.strip()
    )
    ok = pseudo_hits == 0 and len(doc.tables) > 0 and cap_count >= img_count
    return ok, f"tables={len(doc.tables)} pseudo_hits={pseudo_hits} images={img_count} image_captions={cap_count}"


def check_highlights(highlights_path: Path) -> tuple[bool, str]:
    doc = Document(highlights_path)
    items = [p.text.strip().lstrip("•").strip() for p in doc.paragraphs if p.text.strip()]
    lengths = [len(x) for x in items]
    ok = 3 <= len(items) <= 5 and all(n <= 85 for n in lengths)
    return ok, f"count={len(items)} lengths={lengths}"


def run_checks(en_docx: Path, zh_docx: Path, en_pdf: Path, zh_pdf: Path, highlights: Path) -> list[tuple[str, bool, str]]:
    checks: list[tuple[str, bool, str]] = []
    checks.append(("EN file exists", en_docx.exists(), str(en_docx)))
    checks.append(("ZH file exists", zh_docx.exists(), str(zh_docx)))
    checks.append(("EN PDF exists", en_pdf.exists(), str(en_pdf)))
    checks.append(("ZH PDF exists", zh_pdf.exists(), str(zh_pdf)))
    checks.append(("Highlights exists", highlights.exists(), str(highlights)))
    if en_docx.exists():
        checks.append(("EN no Calibri/theme font", *check_fonts_no_calibri(en_docx)))
        checks.append(("EN single-column", *check_single_column(en_docx)))
        checks.append(("EN formula integrity", *check_formula_integrity(en_docx)))
        checks.append(("EN text artifacts", *check_text_artifacts(en_docx)))
        checks.append(("EN abstract/keywords", *check_abstract_keywords(en_docx)))
        checks.append(("EN submission notes/data statement", *check_submission_notes(en_docx)))
        checks.append(("EN tables/captions", *check_tables_and_captions(en_docx)))
    if highlights.exists():
        checks.append(("Highlights constraints", *check_highlights(highlights)))
    return checks


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Validate CEA Word pipeline outputs.")
    parser.add_argument("--en-docx", default="CEA_manuscript_EN.docx")
    parser.add_argument("--zh-docx", default="CEA_manuscript_ZH_working.docx")
    parser.add_argument("--en-pdf", default="CEA_manuscript_EN.pdf")
    parser.add_argument("--zh-pdf", default="CEA_manuscript_ZH_working.pdf")
    parser.add_argument("--highlights", default="highlights.docx")
    parser.add_argument("--report", default="validation_report.txt")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    checks = run_checks(
        Path(args.en_docx),
        Path(args.zh_docx),
        Path(args.en_pdf),
        Path(args.zh_pdf),
        Path(args.highlights),
    )
    lines = []
    fail = 0
    for name, ok, detail in checks:
        status = "PASS" if ok else "FAIL"
        if not ok:
            fail += 1
        lines.append(f"[{status}] {name}: {detail}")
    lines.append(f"\nsummary: total={len(checks)} fail={fail}")
    report_path = Path(args.report)
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print("\n".join(lines))
    return 1 if fail else 0


if __name__ == "__main__":
    raise SystemExit(main())
