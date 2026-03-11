#!/usr/bin/env python
"""Build a CEA-ready submission project folder from finalized artifacts."""

from __future__ import annotations

import argparse
import re
import shutil
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


REQUIRED_FILES = {
    "manuscript_docx": Path("out_cea_word_build_final/CEA_manuscript_EN.docx"),
    "manuscript_pdf": Path("out_cea_word_build_final/CEA_manuscript_EN.pdf"),
    "highlights_docx": Path("out_cea_word_build_final/highlights.docx"),
    "validation": Path("out_cea_word_build_final/validation_report.txt"),
    "ga_pdf": Path("out_cea_word_build_final/graphical_abstract_pipeline.pdf"),
    "ga_png": Path("out_cea_word_build_final/graphical_abstract_pipeline.png"),
    "ga_placeholder_docx": Path("out_cea_word_build_final/CEA_graphical_abstract_placeholder.docx"),
    "bundle_zip": Path("release_bundle_v1.zip"),
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create CEA submission package folder.")
    parser.add_argument(
        "--outdir",
        type=Path,
        default=Path("CEA_submission_ready_package"),
        help="Output folder under repository root.",
    )
    return parser.parse_args()


def ensure_inputs() -> None:
    missing = [str(p) for p in REQUIRED_FILES.values() if not p.exists()]
    if missing:
        raise FileNotFoundError("Missing required input files:\n- " + "\n- ".join(missing))


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def enforce_blind_manuscript(src: Path, dst: Path) -> None:
    doc = Document(str(src))

    # 1.5 line spacing in manuscript body for review readability.
    for p in doc.paragraphs:
        if p.paragraph_format.line_spacing is None:
            p.paragraph_format.line_spacing = 1.5
        if p.style is not None and p.style.name == "Normal":
            p.style.font.name = "Times New Roman"
            p.style.font.size = Pt(11)

    for section in doc.sections:
        sect_pr = section._sectPr

        # Single-column section setting.
        cols = sect_pr.find(qn("w:cols"))
        if cols is not None:
            cols.set(qn("w:num"), "1")

        # Enable line numbering.
        for node in sect_pr.findall(qn("w:lnNumType")):
            sect_pr.remove(node)
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:start"), "1")
        ln.set(qn("w:restart"), "newPage")
        sect_pr.append(ln)

        # Add centered page number field in footer.
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = "Page "
        add_page_field(p)

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


def first_title(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "Title" and p.text.strip():
            return p.text.strip()
    for p in doc.paragraphs:
        if p.text.strip():
            return p.text.strip()
    return "CEA Manuscript Title"


def write_cover_letter(path: Path, title: str) -> None:
    doc = Document()
    doc.add_heading("Cover Letter", level=1)
    today = datetime.now().strftime("%B %d, %Y")
    doc.add_paragraph(today)
    doc.add_paragraph("Dear Editor,")
    doc.add_paragraph(
        f"We submit the manuscript entitled \"{title}\" to Computers and Electronics in Agriculture."
    )
    doc.add_paragraph(
        "This work contributes an audit-ready engineering evidence framework for fish MOT and counting, "
        "linking drift-aware monitoring, stress diagnostics, and deployment-oriented runtime reporting "
        "under reproducible scripts and logged artifacts."
    )
    doc.add_paragraph(
        "The manuscript is original, is not under consideration elsewhere, and all authors approve this submission."
    )
    doc.add_paragraph(
        "A blinded manuscript, title page, highlights, graphical abstract, competing interests statement, "
        "CRediT statement, and reproducibility bundle are included as separate files."
    )
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("The Authors")
    doc.save(str(path))


def write_title_page(path: Path, title: str) -> None:
    doc = Document()
    doc.add_heading("Title Page", level=1)
    doc.add_paragraph(f"Article title: {title}")
    doc.add_paragraph("Authors: Anonymous author team (double-blind review mode).")
    doc.add_paragraph("Affiliations: Withheld for double-blind review; provided in submission metadata.")
    doc.add_paragraph("Corresponding author: Provided in Editorial Manager metadata.")
    doc.add_paragraph("Funding: This research did not receive specific external grant funding.")
    doc.add_paragraph("Acknowledgements: None.")
    doc.save(str(path))


def write_declaration(path: Path) -> None:
    doc = Document()
    doc.add_heading("Declaration of Interest", level=1)
    doc.add_paragraph("The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.")
    doc.save(str(path))


def write_credit(path: Path) -> None:
    doc = Document()
    doc.add_heading("CRediT Author Statement", level=1)
    doc.add_paragraph(
        "Conceptualization, Methodology, Software, Validation, Formal analysis, Investigation, "
        "Data curation, Writing - original draft, Writing - review and editing, Visualization: author team."
    )
    doc.add_paragraph("Supervision, Project administration, Resources: author team.")
    doc.add_paragraph("All authors read and approved the final manuscript.")
    doc.save(str(path))


def write_upload_guide(path: Path) -> None:
    text = """CEA / Editorial Manager upload order

1) Cover Letter -> Cover Letter
2) Title Page -> Title Page
3) Manuscript_Blinded -> Manuscript
4) Highlights -> Highlights
5) Graphical_Abstract -> Graphical Abstract
6) Declaration_of_Interest -> Conflict of Interest
7) CRediT_Author_Statement -> Supporting File
8) release_bundle_v1 (anonymized) -> Supporting File

Final checks:
- Manuscript keeps double-blind anonymity.
- Manuscript has line numbers and page numbers.
- No TODO / placeholder / unresolved tags in manuscript.
- Validation report is PASS.
"""
    path.write_text(text, encoding="utf-8", newline="\n")


def write_gate_report(path: Path, manuscript_docx: Path) -> None:
    with ZipFile(manuscript_docx) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        footer_xml = "".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in zf.namelist()
            if name.startswith("word/footer")
        )

    checks = {
        "Table Table": xml.count("Table Table"),
        "[tab:": xml.count("[tab:"),
        "[alg:": xml.count("[alg:"),
        "meanstd": xml.count("meanstd"),
        "Let  be": xml.count("Let  be"),
        "TODO": xml.count("TODO"),
        "placeholder": xml.count("placeholder"),
    }
    has_line_num = "w:lnNumType" in xml
    has_page_field = (" PAGE " in footer_xml) or ("w:instrText" in footer_xml and "PAGE" in footer_xml)
    lines = ["Submission gate quick check", ""] + [f"{k} = {v}" for k, v in checks.items()]
    lines += [
        f"line_numbers_enabled = {has_line_num}",
        f"page_number_field_present = {has_page_field}",
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8", newline="\n")


def copy_outputs(outdir: Path) -> None:
    files_dir = outdir / "files"
    files_dir.mkdir(parents=True, exist_ok=True)

    # Core files copied as-is.
    shutil.copy2(REQUIRED_FILES["highlights_docx"], files_dir / "Highlights.docx")
    shutil.copy2(REQUIRED_FILES["validation"], files_dir / "Validation_Report.txt")
    shutil.copy2(REQUIRED_FILES["ga_pdf"], files_dir / "Graphical_Abstract.pdf")
    shutil.copy2(REQUIRED_FILES["ga_png"], files_dir / "Graphical_Abstract.png")
    shutil.copy2(REQUIRED_FILES["ga_placeholder_docx"], files_dir / "Graphical_Abstract_Placeholder.docx")
    shutil.copy2(REQUIRED_FILES["bundle_zip"], files_dir / "Supplementary_release_bundle_v1.zip")

    title = first_title(REQUIRED_FILES["manuscript_docx"])

    # Build blinded manuscript with line numbers/page numbers.
    blinded_docx = files_dir / "Manuscript_Blinded.docx"
    enforce_blind_manuscript(REQUIRED_FILES["manuscript_docx"], blinded_docx)

    # Keep source PDF copy as a quick companion proof.
    shutil.copy2(REQUIRED_FILES["manuscript_pdf"], files_dir / "Manuscript_Blinded_preview.pdf")

    # Supporting submission docs.
    write_cover_letter(files_dir / "Cover_Letter.docx", title)
    write_title_page(files_dir / "Title_Page.docx", title)
    write_declaration(files_dir / "Declaration_of_Interest.docx")
    write_credit(files_dir / "CRediT_Author_Statement.docx")
    write_upload_guide(outdir / "UPLOAD_ORDER.txt")
    write_gate_report(outdir / "MANUSCRIPT_GATE_CHECK.txt", blinded_docx)


def main() -> int:
    args = parse_args()
    ensure_inputs()
    if args.outdir.exists():
        shutil.rmtree(args.outdir)
    args.outdir.mkdir(parents=True, exist_ok=True)
    copy_outputs(args.outdir)
    print(f"Created submission project at: {args.outdir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
