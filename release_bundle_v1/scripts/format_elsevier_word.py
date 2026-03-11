from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


TAB_PLACEHOLDER_MAP = {
    "[tab:core_main]": "Table 1",
    "[tab:strong_only]": "Table 2",
    "[tab:all_methods]": "Table 3",
    "[tab:degradation_delta]": "Table 4",
    "[tab:gating_sensitivity]": "Table 5",
    "[tab:count]": "Table 6",
    "[tab:runtime_field_definition]": "Table 7",
    "[tab:runtime_e2e_fields]": "Table 8",
}

FIG_PLACEHOLDER_MAP = {
    "[fig:core_main_png]": "Fig. 1",
    "[fig:strong_png]": "Fig. 2",
    "[fig:all_methods_png]": "Fig. 3",
    "[fig:degradation]": "Fig. 4",
    "[fig:gating_sensitivity]": "Fig. 5",
    "[fig:stratified]": "Fig. 6",
    "[fig:count_png]": "Fig. 7",
    "[fig:runtime]": "Fig. 8",
}

FIGURE_CAPTION_IMAGE_MAP = [
    (
        "Core method family comparison",
        "main_table_val_seedmean_std_paper.png",
    ),
    (
        "Strong baseline comparison",
        "strong_baselines_seedmean_std.png",
    ),
    (
        "Combined comparison of proposed variants and strong baselines",
        "main_table_with_baselines.png",
    ),
    (
        "Controlled degradation curves",
        "degradation_grid.png",
    ),
    (
        "Sensitivity of +gating to threshold values",
        "gating_sensitivity.png",
    ),
    (
        "Stratified metrics over occlusion, density, turning, and low-confidence buckets",
        "stratified_metrics_val.png",
    ),
    (
        "Count stability comparison",
        "count_stability_bar_paper.png",
    ),
    (
        "Runtime profiling",
        "runtime_profile.png",
    ),
]


@dataclass
class FormatStats:
    pseudo_tables_converted: int = 0
    placeholder_refs_fixed: int = 0
    suspicious_formula_todos: int = 0


def write_lines(path: Path | None, lines: list[str]) -> None:
    if path is None:
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    if not lines:
        lines = ["- No entries recorded."]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def paragraph_has_math(paragraph) -> bool:
    return "<m:oMath" in paragraph._p.xml or "<m:oMathPara" in paragraph._p.xml


def set_style_font(style, *, size_pt: float | None, bold: bool | None, italic: bool | None) -> None:
    font = style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0, 0, 0)
    if size_pt is not None:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic

    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), "Times New Roman")


def set_style_paragraph(style, *, line_spacing: float | None = None) -> None:
    pf = style.paragraph_format
    if line_spacing is not None:
        pf.line_spacing = line_spacing


def normalize_styles(doc: DocumentType) -> None:
    styles = doc.styles

    for name in ("Normal", "Body Text", "First Paragraph", "Bibliography", "Abstract"):
        if name in styles:
            set_style_font(styles[name], size_pt=10, bold=False, italic=False)
            set_style_paragraph(styles[name], line_spacing=1.15)

    if "Bibliography" in styles:
        bpf = styles["Bibliography"].paragraph_format
        bpf.left_indent = Inches(0.25)
        bpf.first_line_indent = Inches(-0.25)
        bpf.space_before = Pt(0)
        bpf.space_after = Pt(0)

    if "Title" in styles:
        set_style_font(styles["Title"], size_pt=17, bold=True, italic=False)
        styles["Title"].paragraph_format.alignment = 0  # left

    for name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4", "Abstract Title"):
        if name in styles:
            set_style_font(styles[name], size_pt=10, bold=True, italic=False)
            set_style_paragraph(styles[name], line_spacing=1.15)

    for name in ("Caption", "Captioned Figure", "Image Caption", "Table Caption"):
        if name in styles:
            set_style_font(styles[name], size_pt=9, bold=False, italic=True)
            set_style_paragraph(styles[name], line_spacing=1.0)


def normalize_uncertainty_token(token: str) -> str:
    m = re.fullmatch(r"([+-]?\d+(?:\.\d+)?)\((\d+)\)", token)
    if not m:
        return token

    mean_str, unc_str = m.groups()
    if "." not in mean_str:
        return token

    decimals = len(mean_str.split(".", 1)[1])
    try:
        unc_val = int(unc_str) * (10 ** (-decimals))
    except Exception:
        return token

    unc_fmt = f"{unc_val:.{decimals}f}"
    return f"{mean_str} ± {unc_fmt}"


def normalize_text(text: str) -> str:
    text = text.replace("每", "--")
    text = text.replace("meanstd", "mean±std")
    text = text.replace("¡À", "±")
    text = text.replace("Table Table ", "Table ")
    text = text.replace("Fig. Fig. ", "Fig. ")

    for src, dst in TAB_PLACEHOLDER_MAP.items():
        text = text.replace(src, dst)
    for src, dst in FIG_PLACEHOLDER_MAP.items():
        text = text.replace(src, dst)

    text = re.sub(r"\bdrift-proof\b", "drift-aware monitoring", text, flags=re.IGNORECASE)
    text = re.sub(r"\bdrift proof\b", "drift-aware monitoring", text, flags=re.IGNORECASE)

    def repl_token(match: re.Match[str]) -> str:
        return normalize_uncertainty_token(match.group(0))

    text = re.sub(r"[+-]?\d+(?:\.\d+)?\(\d+\)", repl_token, text)

    # Fix dropped ± caused by inline TeX math rendered as thin-space-separated numbers.
    text = re.sub(r"(\d+\.\d{3})[\u2000-\u200B\u00A0 ]{1,}(\d+\.\d{3})", r"\1 ± \2", text)

    # Repair known formula-placeholder artifacts from TeX->Word conversion.
    if "Let  be the recent input window of size" in text and "reference window" in text:
        text = (
            "Let XtW be the recent input window of size W and Xref be a reference window (calibration period). "
            "Let q(.) denote output-quality statistics computed on a window. "
            "In this manuscript, W denotes the fixed monitoring window length and K denotes the persistence "
            "window count required to trigger an alert."
        )
    text = text.replace("and gated pairs are masked when .", "and gated pairs are masked when s^gate_ij > tau_g.")
    return text


def set_runs_times_new_roman(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{key}"), "Times New Roman")


def replace_plain_paragraph_text(paragraph, new_text: str) -> None:
    # Safe for non-math paragraphs only.
    paragraph.text = new_text


def looks_like_column_spec(text: str) -> bool:
    return bool(re.fullmatch(r"[lcrS|\s]+", text.strip()))


def looks_like_table_line(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if looks_like_column_spec(t):
        return True
    return "&" in t and t.count("&") >= 1


def parse_table_rows(lines: List[str]) -> List[List[str]]:
    rows: List[List[str]] = []
    for line in lines:
        for raw in line.splitlines():
            t = raw.strip()
            if not t:
                continue
            if "Method &" in t:
                t = t[t.index("Method &") :]
            elif "Jitter &" in t:
                t = t[t.index("Jitter &") :]
            elif looks_like_column_spec(t):
                continue
            if "&" not in t:
                continue
            t = t.rstrip("\\").strip()
            parts = [normalize_text(x.strip()) for x in t.split("&")]
            rows.append(parts)
    return rows


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def apply_elsevier_table_look(table) -> None:
    # remove all vertical borders; keep top, header-bottom, bottom lines.
    rows = table.rows
    cols = len(rows[0].cells) if rows else 0
    if cols == 0:
        return

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = row.cells[c_idx]
            set_cell_border(
                cell,
                left={"val": "nil"},
                right={"val": "nil"},
                top={"val": "nil"},
                bottom={"val": "nil"},
            )

    # top line
    for cell in rows[0].cells:
        set_cell_border(cell, top={"val": "single", "sz": "8", "color": "000000"})
    # header separator
    for cell in rows[0].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": "000000"})
    # bottom line
    for cell in rows[-1].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": "000000"})

    for cell in rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def style_all_tables(doc: DocumentType) -> int:
    styled = 0
    for table in doc.tables:
        apply_elsevier_table_look(table)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    para.paragraph_format.line_spacing = 1.0
        styled += 1
    return styled


def convert_pseudo_table_blocks(doc: DocumentType) -> int:
    converted = 0
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        if not looks_like_table_line(para.text):
            i += 1
            continue

        start = i
        lines = []
        while i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            if not looks_like_table_line(p.text):
                break
            lines.append(p.text)
            i += 1

        rows = parse_table_rows(lines)
        if len(rows) < 2:
            continue

        ncols = max(len(r) for r in rows)
        first_para = doc.paragraphs[start]
        table = doc.add_table(rows=len(rows), cols=ncols)
        for r_idx, row_vals in enumerate(rows):
            for c_idx in range(ncols):
                val = row_vals[c_idx] if c_idx < len(row_vals) else ""
                table.cell(r_idx, c_idx).text = val

        apply_elsevier_table_look(table)

        first_para._p.addprevious(table._tbl)

        for idx in range(start + len(lines) - 1, start - 1, -1):
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)

        converted += 1
        i = start + 1

    return converted


def normalize_paragraphs(doc: DocumentType, stats: FormatStats) -> None:
    for para in doc.paragraphs:
        if paragraph_has_math(para):
            # keep math paragraphs intact but still normalize visible text runs where present
            for run in para.runs:
                if run.text:
                    run.text = normalize_text(run.text)
            set_runs_times_new_roman(para)
            continue

        old = para.text
        new = normalize_text(old)

        if "Submission notes." in old:
            new = (
                "Submission notes. Highlights: submitted as a separate file. "
                "Graphical abstract: submitted as a separate file (see graphical abstract guidelines). "
                "Graphical Abstract - Figure GA (submitted separately)."
            )

        if old != new:
            replace_plain_paragraph_text(para, new)
            if "Table " in new or "Fig. " in new:
                stats.placeholder_refs_fixed += 1

        set_runs_times_new_roman(para)


def ensure_heading_alignment(doc: DocumentType) -> None:
    for para in doc.paragraphs:
        if para.style and para.style.name in {"Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"}:
            para.paragraph_format.alignment = 0


def ensure_submission_note(doc: DocumentType) -> None:
    target = (
        "Submission notes. Highlights: submitted as a separate file. "
        "Graphical abstract: submitted as a separate file (see graphical abstract guidelines). "
        "Graphical Abstract - Figure GA (submitted separately)."
    )
    for para in doc.paragraphs:
        if "Submission notes." in para.text:
            para.text = target
            set_runs_times_new_roman(para)
            return

    insert_idx = 0
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == "abstract":
            insert_idx = i + 2
            break
    if insert_idx < len(doc.paragraphs):
        p = doc.paragraphs[insert_idx].insert_paragraph_before(target)
    else:
        p = doc.add_paragraph(target)
    if "Body Text" in doc.styles:
        p.style = doc.styles["Body Text"]
    set_runs_times_new_roman(p)


def ensure_data_statement(doc: DocumentType, todos: list[str], changes: list[str]) -> None:
    full = "\n".join(p.text for p in doc.paragraphs)
    if "Data availability" in full or "Data availability." in full:
        return
    h = doc.add_paragraph("Data availability and reproducibility statement")
    if "Heading 1" in doc.styles:
        h.style = doc.styles["Heading 1"]
    p = doc.add_paragraph(
        "Data availability. TODO: add data sharing scope/restriction statement and derived artifact access details."
    )
    if "Body Text" in doc.styles:
        p.style = doc.styles["Body Text"]
    set_runs_times_new_roman(p)
    todos.append("TODO-DATA: Added placeholder Data availability statement; complete with final access policy.")
    changes.append("Inserted missing Data availability statement placeholder.")


def enforce_cea_front_matter(doc: DocumentType, todos: list[str], changes: list[str]) -> None:
    ensure_submission_note(doc)
    ensure_data_statement(doc, todos, changes)

    abstract_words = None
    keyword_count = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == "abstract":
            for j in range(i + 1, min(i + 8, len(doc.paragraphs))):
                body = doc.paragraphs[j]
                if not body.text.strip():
                    continue
                abstract_words = len(re.findall(r"[A-Za-z0-9\-\+]+", body.text))
                if abstract_words > 250 and "TODO-ABS1:" not in body.text:
                    body.text = body.text.strip() + " TODO-ABS1: shorten abstract to <=250 words before submission."
                    todos.append("TODO-ABS1: Abstract exceeds 250 words; shorten before submission.")
                break
            if i + 2 < len(doc.paragraphs):
                kline = doc.paragraphs[i + 2].text.replace("\\sep", ",")
                keys = [k.strip() for k in re.split(r",|;", kline) if k.strip()]
                keyword_count = len(keys)
                if not (1 <= keyword_count <= 7):
                    t = doc.paragraphs[i + 2].insert_paragraph_before(
                        "TODO-KEY1: Ensure 1-7 keywords in final manuscript."
                    )
                    if "Body Text" in doc.styles:
                        t.style = doc.styles["Body Text"]
                    todos.append("TODO-KEY1: Keyword count not in 1-7 range.")
            break

    changes.append(
        f"Front-matter compliance checked (abstract_words={abstract_words}, keyword_count={keyword_count})."
    )


def enforce_content_safety_and_scope(doc: DocumentType, todos: list[str], changes: list[str]) -> None:
    full = "\n".join(p.text.lower() for p in doc.paragraphs)
    if "drift-proof" in full:
        todos.append("TODO-TERM1: Residual 'drift-proof' text detected; replace with drift-aware phrasing.")
    if "terminology and scope" not in full:
        h = doc.add_paragraph("Terminology and scope")
        if "Heading 2" in doc.styles:
            h.style = doc.styles["Heading 2"]
        p = doc.add_paragraph(
            "This study performs drift-aware monitoring and response only; it does not guarantee drift elimination."
        )
        if "Body Text" in doc.styles:
            p.style = doc.styles["Body Text"]
        todos.append("TODO-SCOPE1: Verify inserted Terminology and scope section position.")
        changes.append("Inserted missing Terminology and scope section placeholder.")

    if "scope a" not in full or "scope b" not in full:
        todos.append("TODO-RUNTIME1: Scope A/B wording missing; verify runtime scope split.")
    if "gating activation analysis" not in full:
        todos.append("TODO-GATE1: Add gating activation diagnostics subsection.")
    changes.append("Content-safety checks applied (drift-aware wording, scope/gating presence).")


def parse_pandoc_math_warnings(pandoc_log_path: Path | None) -> list[str]:
    if pandoc_log_path is None or not pandoc_log_path.exists():
        return []
    warnings: list[str] = []
    for line in pandoc_log_path.read_text(encoding="utf-8", errors="replace").splitlines():
        if "Could not convert TeX math" in line:
            warnings.append(line.strip())
    return warnings


def fix_missing_equations(
    doc: DocumentType,
    pandoc_log_path: Path | None,
    todos: list[str],
    changes: list[str],
    stats: FormatStats,
) -> None:
    warnings = parse_pandoc_math_warnings(pandoc_log_path)
    for w in warnings:
        p = doc.add_paragraph(f"TODO-MATH: restore equation from pandoc warning: {w}")
        if "Body Text" in doc.styles:
            p.style = doc.styles["Body Text"]
        todos.append(f"TODO-MATH: {w}")

    # Existing in-doc blank-formula heuristic.
    check_suspicious_formula_blanks(doc, stats)
    if warnings:
        changes.append(f"Registered {len(warnings)} math-conversion TODOs from pandoc log.")
    if stats.suspicious_formula_todos:
        todos.append(
            f"TODO-MATH-BLANK: {stats.suspicious_formula_todos} blank-formula contexts flagged in document text."
        )

def normalize_reference_text(text: str) -> str:
    t = text
    t = t.replace("“", "").replace("”", "")
    t = t.replace("‘", "").replace("’", "'")
    t = re.sub(r"\s+", " ", t).strip()
    t = re.sub(r"\s+\.", ".", t)
    t = re.sub(r"\s+,", ",", t)
    return t


def format_references_section(doc: DocumentType) -> int:
    ref_paras = [p for p in doc.paragraphs if p.style and p.style.name == "Bibliography" and p.text.strip()]
    if not ref_paras:
        return 0

    first_para = ref_paras[0]

    count = 0
    for p in ref_paras:
        new_text = normalize_reference_text(p.text)
        if new_text != p.text:
            p.text = new_text
        if "Bibliography" in doc.styles:
            p.style = doc.styles["Bibliography"]
        pf = p.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.25)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.15
        p.paragraph_format.alignment = 0
        set_runs_times_new_roman(p)
        count += 1

    ref_start_idx = 0
    for idx, p in enumerate(doc.paragraphs):
        if p._p is first_para._p:
            ref_start_idx = idx
            break
    has_heading = ref_start_idx > 0 and doc.paragraphs[ref_start_idx - 1].text.strip().lower() == "references"
    if not has_heading:
        h = first_para.insert_paragraph_before("References")
        if "Heading 1" in doc.styles:
            h.style = doc.styles["Heading 1"]
        for run in h.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    return count


def paragraph_has_drawing(paragraph) -> bool:
    return "<w:drawing" in paragraph._p.xml


def pick_result_figure(caption_text: str, figs_dir: Path) -> Path | None:
    t = " ".join(caption_text.strip().split()).lower()
    for key, filename in FIGURE_CAPTION_IMAGE_MAP:
        if key.lower() in t:
            candidate = figs_dir / filename
            if candidate.exists():
                return candidate
    return None


def insert_result_figures(doc: DocumentType, figs_dir: Path) -> int:
    inserted = 0
    paragraphs = list(doc.paragraphs)
    for i, para in enumerate(paragraphs):
        if not (para.style and para.style.name == "Image Caption"):
            continue

        image_path = pick_result_figure(para.text, figs_dir)
        if image_path is None:
            continue

        if paragraph_has_drawing(para):
            continue

        target = None
        if i > 0:
            prev = paragraphs[i - 1]
            if paragraph_has_drawing(prev):
                continue
            if prev.style and prev.style.name == "Captioned Figure":
                target = prev

        if target is None:
            target = para.insert_paragraph_before("")
            if "Captioned Figure" in doc.styles:
                target.style = doc.styles["Captioned Figure"]
            elif "Body Text" in doc.styles:
                target.style = doc.styles["Body Text"]

        run = target.add_run()
        run.add_picture(str(image_path), width=Inches(6.0))
        target.paragraph_format.alignment = 1  # center
        inserted += 1
    return inserted


def ensure_abstract_word_limit(doc: DocumentType) -> None:
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == "Abstract Title":
            # next non-empty paragraph is abstract body
            for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
                p = doc.paragraphs[j]
                if p.text.strip():
                    words = re.findall(r"[A-Za-z0-9\-\+]+", p.text)
                    if len(words) > 250:
                        p.text = p.text.strip() + " TODO-ABS1: shorten abstract to <=250 words before submission."
                    return


def check_suspicious_formula_blanks(doc: DocumentType, stats: FormatStats) -> None:
    # Add TODO only for suspicious blank equation contexts that have no OMML.
    for para in doc.paragraphs:
        if paragraph_has_math(para):
            continue
        t = re.sub(r"\s+", " ", para.text).strip()
        if re.search(r"\bLet\s+be\b", t) or re.search(r"\bis\s+unitless\b", t) and t.startswith("is "):
            todo = (
                "TODO-FORMULA: Potential missing equation detected in this sentence. "
                "Cross-check with main.tex drift-equation lines and restore as editable OMML equation."
            )
            para.add_run(" " + todo)
            stats.suspicious_formula_todos += 1


def create_highlights_docx(path: Path) -> None:
    bullets = [
        "Audit-ready protocol links fairness checks, stress tests, and deployment rules.",
        "Drift-aware diagnostics localize failures by bucket and activation behavior.",
        "Candidate threshold intervals are evaluated via activation-level diagnostics.",
        "Scope A/B runtime reporting separates tracker cost from end-to-end cost.",
    ]
    for line in bullets:
        if len(line) > 85:
            raise ValueError(f"Highlight exceeds 85 chars: {line}")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    for line in bullets:
        p = doc.add_paragraph(line, style="List Bullet")
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)

    doc.save(path)


def create_graphical_abstract_placeholder(path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("Graphical Abstract (submitted separately)")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
    doc.save(path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Format manuscript docx in Elsevier-like style.")
    parser.add_argument(
        "--input",
        "--in",
        dest="input",
        default="paper/cea_draft/CEA_paper_formatted_like_elsevier_v3.docx",
        help="Input DOCX manuscript.",
    )
    parser.add_argument(
        "--output",
        "--out",
        dest="output",
        default="paper/cea_draft/CEA_manuscript_elsevier_style.docx",
        help="Output formatted DOCX manuscript.",
    )
    parser.add_argument(
        "--highlights",
        default="paper/cea_draft/CEA_highlights.docx",
        help="Output Highlights DOCX.",
    )
    parser.add_argument(
        "--ga-placeholder",
        default="paper/cea_draft/CEA_graphical_abstract_placeholder.docx",
        help="Output Graphical Abstract placeholder DOCX.",
    )
    parser.add_argument(
        "--todo",
        default="",
        help="Optional TODO markdown output path.",
    )
    parser.add_argument(
        "--changes",
        default="",
        help="Optional changes markdown output path.",
    )
    parser.add_argument(
        "--pandoc-log",
        default="",
        help="Optional pandoc stderr log path (for math warning TODO extraction).",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    in_path = Path(args.input)
    out_path = Path(args.output)
    highlights_path = Path(args.highlights)
    ga_path = Path(args.ga_placeholder)
    todo_path = Path(args.todo) if args.todo else None
    changes_path = Path(args.changes) if args.changes else None
    pandoc_log_path = Path(args.pandoc_log) if args.pandoc_log else None

    doc = Document(in_path)
    stats = FormatStats()
    changes: list[str] = []
    todos: list[str] = []
    repo_root = Path(__file__).resolve().parents[1]
    figs_dir = repo_root / "paper" / "cea_draft" / "figs"

    changes.append("Loaded input DOCX.")
    normalize_styles(doc)
    changes.append("Applied style contract (Times New Roman, heading/caption sizes, spacing).")
    normalize_paragraphs(doc, stats)
    enforce_content_safety_and_scope(doc, todos, changes)
    enforce_cea_front_matter(doc, todos, changes)
    stats.pseudo_tables_converted = convert_pseudo_table_blocks(doc)
    changes.append(f"Converted pseudo tables: {stats.pseudo_tables_converted}.")
    styled_tables = style_all_tables(doc)
    changes.append(f"Applied table styling to {styled_tables} tables.")
    inserted_figures = insert_result_figures(doc, figs_dir)
    changes.append(f"Inserted/confirmed result figures: {inserted_figures}.")
    ensure_submission_note(doc)
    refs_formatted = format_references_section(doc)
    changes.append(f"Formatted references: {refs_formatted} entries.")
    ensure_heading_alignment(doc)
    ensure_abstract_word_limit(doc)
    fix_missing_equations(doc, pandoc_log_path, todos, changes, stats)
    changes.append("Normalized symbols/cross-references/special characters.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    create_highlights_docx(highlights_path)
    create_graphical_abstract_placeholder(ga_path)
    changes.append("Generated highlights and graphical abstract placeholder files.")
    write_lines(changes_path, [f"- {x}" for x in changes])
    write_lines(todo_path, [f"- {x}" for x in todos])

    print(f"Saved manuscript: {out_path}")
    print(f"Saved highlights: {highlights_path}")
    print(f"Saved graphical abstract placeholder: {ga_path}")
    if changes_path is not None:
        print(f"Saved changes: {changes_path}")
    if todo_path is not None:
        print(f"Saved todos: {todo_path}")
    print(
        "stats "
        f"pseudo_tables_converted={stats.pseudo_tables_converted} "
        f"placeholder_refs_fixed={stats.placeholder_refs_fixed} "
        f"suspicious_formula_todos={stats.suspicious_formula_todos} "
        f"tables_styled={styled_tables} "
        f"figures_inserted={inserted_figures} "
        f"references_formatted={refs_formatted}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
